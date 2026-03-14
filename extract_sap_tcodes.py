"""
SAP Transaction Code Extractor
================================
Scans a folder of Word (.docx) documents, extracts SAP transaction codes,
and outputs an Excel file with:
  Column A: Transaction Code  (distinct - merged cell)
  Column B: Document Name     (multiple rows per code)
  Column C: Source Sentence   (the line where the code was found)

Logic: A transaction code is ONLY extracted when it immediately follows
a known SAP trigger phrase (Run, T-code, Tcode, SAP Transaction, etc.)
exactly as shown in the reference screenshot.

Usage: place .docx files in the input/ folder, run script, get output/SAP_Tcodes_Extract.xlsx
"""

import os
import re
from collections import defaultdict

try:
    from docx import Document
except ImportError:
    raise SystemExit("python-docx not installed. Run: pip install python-docx")

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
except ImportError:
    raise SystemExit("openpyxl not installed. Run: pip install openpyxl")


# ---------------------------------------------------------------------------
# EXCLUSION LIST — add any transaction codes you want to exclude from output.
# Examples: EXCLUDE_TCODES = {'SE16', 'SM30', 'FB01'}
# Case-insensitive. Codes in this list will not appear in any sheet.
# ---------------------------------------------------------------------------
EXCLUDE_TCODES = set()   # <-- add codes here, e.g. {'SE16', 'SM30'}


# ---------------------------------------------------------------------------
# Trigger-based extraction patterns
# Each pattern captures the transaction code in group(1).
# Ordered from most specific to least specific.
# ---------------------------------------------------------------------------
TRIGGER_PATTERNS = [
    # "Run SAP t code CJ88" / "Run SAP t-code CJ88"
    re.compile(r'Run\s+SAP\s+[Tt][-\s]?[Cc]ode\s+([A-Z][A-Z0-9_\-]{1,29})', re.IGNORECASE),

    # "Run the YRTR_ASSET_BALANCES Tcode"  (code comes BEFORE the word Tcode)
    re.compile(r'Run\s+the\s+([A-Z][A-Z0-9_\-]{1,29})\s+[Tt]code', re.IGNORECASE),

    # "Run S_ALR-87013611 SAP Tcode" (code before trailing SAP Tcode)
    re.compile(r'Run\s+([A-Z][A-Z0-9_\-]{1,29})\s+SAP\s+[Tt]code', re.IGNORECASE),

    # "Run S_ALR_87011964" / "Run FAGLB03" — only when followed by SAP context words
    # Postfix: SAP, Tcode, T-code, T code, Transaction (case-insensitive)
    re.compile(r'Run\s+([A-Z][A-Z0-9_\-]{1,29})\s+(?:SAP|[Tt][-\s]?[Cc]ode|[Tt]ransaction)\b', re.IGNORECASE),

    # "T-code ZRTR_GL_PARK_GL" / "T-code (CJ20N)" / "Tcode CJ20N" / "Tcode: FB01"
    re.compile(r'[Tt][-\s]?[Cc]ode[s]?\s*[:\(]?\s*([A-Z][A-Z0-9_\-]{1,29})\)?', re.IGNORECASE),

    # "SAP Transaction CJ20N" / "SAP Transaction (CJ20N)" — SAP directly before Transaction
    re.compile(r'SAP\s+[Tt]ransaction\s*[:\(]?\s*([A-Z][A-Z0-9_\-]{1,29})\)?', re.IGNORECASE),

    # "SAP Asset transaction (S_ALR_87012048)" — SAP + any words + transaction + (CODE)
    re.compile(r'SAP\s+\w[\w\s]{0,40}[Tt]ransaction\s*\(\s*([A-Z][A-Z0-9_\-]{1,29})\s*\)', re.IGNORECASE),

    # "Transaction code FB01" / "Transaction: FB01"
    re.compile(r'[Tt]ransaction\s+[Cc]ode\s*[:\(]?\s*([A-Z][A-Z0-9_\-]{1,29})\)?', re.IGNORECASE),
]

# ---------------------------------------------------------------------------
# Table-based extraction: column headers that signal SAP code columns
# ---------------------------------------------------------------------------
TABLE_CODE_HEADERS = {
    'process terms', 'process acronyms', 'process term', 'process acronym',
    'transaction code', 'transaction codes', 't-code', 'tcode', 't code',
    'sap transaction', 'sap tcode', 'sap t-code', 'sap code', 'sap codes',
}

# Cell values inside SAP-code tables that should be ignored (not real codes)
TABLE_IGNORE_VALUES = {
    'NA', 'N/A', 'N.A', 'N.A.', 'NONE', 'NIL', 'TBD', 'TBC', '-', '--', '---',
    'WBS', 'WNS', 'N', 'A', 'YES', 'NO', 'TRUE', 'FALSE',
}

# Words that must never be treated as transaction codes
NOT_A_TCODE = {
    'THE', 'AND', 'FOR', 'RUN', 'SAP', 'CODE', 'CODES', 'WITH', 'FROM', 'INTO',
    'THIS', 'THAT', 'HAVE', 'BEEN', 'WILL', 'ALSO', 'CAN', 'NOT', 'ARE', 'ALL',
    'BUT', 'USE', 'NEW', 'OLD', 'END', 'ADD', 'SET', 'GET', 'PUT', 'OUT', 'OFF',
    'TOP', 'BOX', 'YES', 'NO', 'OK', 'GO', 'DO', 'IF', 'AS', 'AT', 'BE', 'BY',
    'IN', 'IS', 'IT', 'OF', 'ON', 'OR', 'SO', 'TO', 'UP', 'US', 'WE', 'ME',
    'TCODE', 'TCODES', 'TRANSACTION', 'REPORT', 'MODULE', 'SYSTEM', 'TABLE',
    'FIELD', 'VALUE', 'PLEASE', 'CLICK', 'OPEN', 'CLOSE', 'ENTER', 'SELECT',
    'PRESS', 'BUTTON', 'SCREEN', 'WINDOW', 'MENU', 'LIST', 'VIEW', 'NEXT',
    'BACK', 'SAVE', 'EXIT', 'NOTE', 'STEP', 'THEN', 'WHEN', 'ONCE', 'AFTER',
    'BEFORE', 'USING', 'BELOW', 'ABOVE', 'RIGHT', 'LEFT', 'HERE', 'THERE',
}


def is_valid_tcode(candidate: str) -> bool:
    c = candidate.upper().strip().rstrip('.,;)(')
    if not c or c in NOT_A_TCODE:
        return False
    if len(c) < 2 or not c[0].isalpha():
        return False
    # Only alphanumeric + underscore
    if not re.match(r'^[A-Z][A-Z0-9_]+$', c):
        return False
    return True


def extract_tcodes_from_text(text: str) -> list:
    """Return list of (tcode, source_sentence) — trigger-phrase based only."""
    results = []
    text_stripped = text.strip()
    if not text_stripped:
        return results

    seen = set()
    for pattern in TRIGGER_PATTERNS:
        for m in pattern.finditer(text):
            raw = m.group(1).strip().rstrip('.,;)(')
            candidate = raw.upper()
            if is_valid_tcode(candidate) and candidate not in seen:
                seen.add(candidate)
                candidate = candidate.replace('-', '_')
                results.append((candidate, text_stripped))

    return results


def is_sap_table_header(text: str) -> bool:
    """Return True if this cell text is a recognised SAP-code table header."""
    return text.strip().lower() in TABLE_CODE_HEADERS


def extract_from_table(table) -> list:
    """
    Scan a table for SAP-code columns.
    If a column header matches TABLE_CODE_HEADERS, treat every non-header
    cell in that column as a transaction code (after filtering ignore values).
    Also runs trigger-phrase extraction on every cell as normal.
    """
    matches = []
    if not table.rows:
        return matches

    rows = list(table.rows)

    # --- Pass 1: find which columns are SAP-code columns by checking row 0 ---
    sap_col_indices = set()
    header_row = rows[0]
    for col_idx, cell in enumerate(header_row.cells):
        cell_text = cell.text.strip()
        if is_sap_table_header(cell_text):
            sap_col_indices.add(col_idx)

    # --- Pass 2: extract codes from SAP-code columns (skip header row) ---
    for row_idx, row in enumerate(rows):
        for col_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            if not cell_text:
                continue

            if col_idx in sap_col_indices and row_idx > 0:
                # This cell is under a SAP-code header — treat value as code directly
                candidate = cell_text.upper().replace('-', '_')
                if candidate not in TABLE_IGNORE_VALUES and is_valid_tcode(candidate):
                    # Source sentence = "Process Terms: <value>" for clarity
                    header_label = header_row.cells[col_idx].text.strip()
                    sentence = f"{header_label}: {cell_text}"
                    matches.append((candidate, sentence))
            else:
                # Normal trigger-phrase extraction on every cell
                for tcode, sentence in extract_tcodes_from_text(cell_text):
                    matches.append((tcode, sentence))

    return matches


def extract_from_docx(filepath: str) -> list:
    """Return list of (tcode, source_sentence) for all matches in a docx."""
    doc = Document(filepath)
    matches = []

    for para in doc.paragraphs:
        for tcode, sentence in extract_tcodes_from_text(para.text):
            matches.append((tcode, sentence))

    for table in doc.tables:
        matches.extend(extract_from_table(table))

    return matches


def scan_folder(folder: str):
    """
    Recursively scan all .docx files in folder and subfolders.
    Returns:
      tcode_data : {tcode: [(rel_path, sentence), ...]}
      others     : [rel_paths with no t-codes found]
    """
    tcode_data = defaultdict(list)
    others = []

    all_docx = []
    for dirpath, _, filenames in os.walk(folder):
        for f in filenames:
            if f.lower().endswith('.docx') and not f.startswith('~$'):
                all_docx.append(os.path.join(dirpath, f))
    all_docx.sort()

    if not all_docx:
        raise SystemExit(
            f"No .docx files found in '{folder}' or any of its subfolders.\n"
            f"Place your Word documents inside the 'input' folder and try again."
        )

    print(f"Found {len(all_docx)} Word document(s) to scan...\n")

    for filepath in all_docx:
        rel_path = os.path.relpath(filepath, folder)
        try:
            matches = extract_from_docx(filepath)
        except Exception as e:
            print(f"  [WARN] Could not read '{rel_path}': {e}")
            others.append(rel_path)
            continue

        if matches:
            for tcode, sentence in matches:
                tcode_data[tcode].append((rel_path, sentence))
            unique = len(set(m[0] for m in matches))
            print(f"  + {rel_path} -> {unique} t-code(s) found")
        else:
            others.append(rel_path)
            print(f"  - {rel_path} -> no t-codes found (-> Others)")

    # Apply exclusion list — remove any codes in EXCLUDE_TCODES
    excluded = {t.upper() for t in EXCLUDE_TCODES}
    tcode_data = {k: v for k, v in tcode_data.items() if k.upper() not in excluded}

    return tcode_data, others


def apply_header(ws, headers, hdr_font, hdr_fill, hdr_align, medium_border):
    for col, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=col, value=h)
        c.font = hdr_font; c.fill = hdr_fill
        c.alignment = hdr_align; c.border = medium_border
    ws.row_dimensions[1].height = 30


def build_sheet2(wb, tcode_data, others, styles):
    """Sheet 2: Transaction Code | Document Name  (one row per unique tcode+doc pair)."""
    ws = wb.create_sheet("T-Code vs Document")
    hdr_font, hdr_fill, hdr_align, medium_border, thin_border = styles
    doc_font  = Font(name='Arial', size=10)
    tcode_font = Font(name='Courier New', bold=True, color='1F4E79', size=10)
    top_align  = Alignment(horizontal='left', vertical='top', wrap_text=True)
    mid_align  = Alignment(horizontal='center', vertical='center', wrap_text=True)
    alt_fill   = PatternFill('solid', start_color='EBF3FB')
    base_fill  = PatternFill('solid', start_color='FFFFFF')

    apply_header(ws, ['Transaction Code', 'Document Name'], hdr_font, hdr_fill, hdr_align, medium_border)

    row = 2
    block = 0
    for tcode in sorted(tcode_data.keys()):
        # Unique documents for this tcode (preserve order, deduplicate)
        seen_docs = []
        seen_set  = set()
        for rel_path, _ in tcode_data[tcode]:
            if rel_path not in seen_set:
                seen_docs.append(rel_path)
                seen_set.add(rel_path)

        fill = alt_fill if (block % 2 == 0) else base_fill
        block += 1
        start_row = row

        for doc in seen_docs:
            c = ws.cell(row=row, column=2, value=doc)
            c.font = doc_font; c.fill = fill
            c.alignment = top_align; c.border = thin_border
            ws.row_dimensions[row].height = 28
            row += 1

        # Tcode in col A — merged across its doc rows
        c = ws.cell(row=start_row, column=1, value=tcode)
        c.font = tcode_font; c.fill = fill
        c.alignment = mid_align; c.border = medium_border
        if len(seen_docs) > 1:
            ws.merge_cells(start_row=start_row, start_column=1,
                           end_row=row - 1, end_column=1)
            c = ws.cell(row=start_row, column=1)
            c.font = tcode_font; c.fill = fill
            c.alignment = mid_align; c.border = medium_border

    # Others rows
    if others:
        others_fill = PatternFill('solid', start_color='FFE7E7')
        others_font_s2 = Font(name='Courier New', bold=True, color='C00000', size=10)
        start_row = row
        n = len(others)
        for rel_path in sorted(others):
            fill = alt_fill if (row % 2 == 0) else base_fill
            c = ws.cell(row=row, column=2, value=rel_path)
            c.font = doc_font; c.fill = fill
            c.alignment = top_align; c.border = thin_border
            ws.row_dimensions[row].height = 28
            row += 1
        c = ws.cell(row=start_row, column=1, value='Others')
        c.font = others_font_s2; c.fill = others_fill
        c.alignment = mid_align; c.border = medium_border
        if n > 1:
            ws.merge_cells(start_row=start_row, start_column=1,
                           end_row=row - 1, end_column=1)
            c = ws.cell(row=start_row, column=1)
            c.font = others_font_s2; c.fill = others_fill
            c.alignment = mid_align; c.border = medium_border

    ws.column_dimensions['A'].width = 26
    ws.column_dimensions['B'].width = 55
    ws.freeze_panes = 'B2'


def build_sheet3(wb, tcode_data, others, styles):
    """Sheet 3: Transaction Code | Count of Documents  (sorted by count desc)."""
    ws = wb.create_sheet("T-Code Count")
    hdr_font, hdr_fill, hdr_align, medium_border, thin_border = styles
    tcode_font  = Font(name='Courier New', bold=True, color='1F4E79', size=10)
    count_font  = Font(name='Arial', bold=True, size=10)
    top_align   = Alignment(horizontal='left',   vertical='center', wrap_text=False)
    cnt_align   = Alignment(horizontal='center', vertical='center')
    alt_fill    = PatternFill('solid', start_color='EBF3FB')
    base_fill   = PatternFill('solid', start_color='FFFFFF')

    apply_header(ws, ['Transaction Code', 'Document Count'], hdr_font, hdr_fill, hdr_align, medium_border)

    # Count unique docs per tcode, sort by count descending
    counts = []
    for tcode, entries in tcode_data.items():
        unique_docs = len({rel_path for rel_path, _ in entries})
        counts.append((tcode, unique_docs))
    counts.sort(key=lambda x: (-x[1], x[0]))

    for i, (tcode, cnt) in enumerate(counts):
        row  = i + 2
        fill = alt_fill if (i % 2 == 0) else base_fill

        c = ws.cell(row=row, column=1, value=tcode)
        c.font = tcode_font; c.fill = fill
        c.alignment = top_align; c.border = thin_border

        c = ws.cell(row=row, column=2, value=cnt)
        c.font = count_font; c.fill = fill
        c.alignment = cnt_align; c.border = thin_border
        ws.row_dimensions[row].height = 24

    # Others row — count of unique docs in others list
    if others:
        i = len(counts)
        row = i + 2
        fill = alt_fill if (i % 2 == 0) else base_fill
        others_font_s3 = Font(name='Courier New', bold=True, color='C00000', size=10)
        c = ws.cell(row=row, column=1, value='Others')
        c.font = others_font_s3; c.fill = PatternFill('solid', start_color='FFE7E7')
        c.alignment = top_align; c.border = thin_border
        c = ws.cell(row=row, column=2, value=len(others))
        c.font = count_font; c.fill = PatternFill('solid', start_color='FFE7E7')
        c.alignment = cnt_align; c.border = thin_border
        ws.row_dimensions[row].height = 24

    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 20
    ws.freeze_panes = 'A2'


def build_sheet4(wb, tcode_data, others, input_dir, styles):
    """Sheet 4: Folder Name | Document Name | Transaction Codes (one row per doc)."""
    ws = wb.create_sheet("Doc vs T-Codes")
    hdr_font, hdr_fill, hdr_align, medium_border, thin_border = styles
    folder_font  = Font(name='Arial', bold=True, color='1F4E79', size=10)
    doc_font     = Font(name='Arial', size=10)
    tcode_font   = Font(name='Courier New', size=10)
    others_font  = Font(name='Courier New', bold=True, color='C00000', size=10)
    top_align    = Alignment(horizontal='left', vertical='top', wrap_text=True)
    mid_align    = Alignment(horizontal='center', vertical='center', wrap_text=True)
    alt_fill     = PatternFill('solid', start_color='EBF3FB')
    base_fill    = PatternFill('solid', start_color='FFFFFF')
    others_fill  = PatternFill('solid', start_color='FFE7E7')

    apply_header(ws, ['Folder Name', 'Document Name', 'Transaction Codes'],
                 hdr_font, hdr_fill, hdr_align, medium_border)

    # Build doc -> set of tcodes mapping
    doc_to_tcodes = {}
    for tcode, entries in tcode_data.items():
        for rel_path, _ in entries:
            doc_to_tcodes.setdefault(rel_path, set()).add(tcode)

    # Also add others docs (no tcodes)
    for rel_path in others:
        doc_to_tcodes.setdefault(rel_path, set())

    row = 2
    block = 0
    # Group docs by folder
    folder_map = {}
    for rel_path in sorted(doc_to_tcodes.keys()):
        parts = rel_path.replace('\\', '/').split('/')
        folder = '/'.join(parts[:-1]) if len(parts) > 1 else '(root)'
        doc    = parts[-1]
        folder_map.setdefault(folder, []).append((doc, rel_path))

    for folder in sorted(folder_map.keys()):
        docs     = folder_map[folder]
        n        = len(docs)
        start_row = row
        fill     = alt_fill if (block % 2 == 0) else base_fill
        block   += 1

        for doc_name, rel_path in docs:
            tcodes = sorted(doc_to_tcodes.get(rel_path, set()))
            is_other = not tcodes

            tcode_str = ', '.join(tcodes) if tcodes else 'No SAP transaction code detected'

            c = ws.cell(row=row, column=2, value=doc_name)
            c.font = doc_font; c.fill = others_fill if is_other else fill
            c.alignment = top_align; c.border = thin_border

            c = ws.cell(row=row, column=3, value=tcode_str)
            c.font = others_font if is_other else tcode_font
            c.fill = others_fill if is_other else fill
            c.alignment = top_align; c.border = thin_border

            ws.row_dimensions[row].height = 35
            row += 1

        # Folder name in col A — merged across its doc rows
        c = ws.cell(row=start_row, column=1, value=folder)
        c.font = folder_font; c.fill = fill
        c.alignment = mid_align; c.border = medium_border
        if n > 1:
            ws.merge_cells(start_row=start_row, start_column=1,
                           end_row=row - 1, end_column=1)
            c = ws.cell(row=start_row, column=1)
            c.font = folder_font; c.fill = fill
            c.alignment = mid_align; c.border = medium_border

    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 40
    ws.column_dimensions['C'].width = 70
    ws.freeze_panes = 'B2'


def build_excel(tcode_data: dict, others: list, output_path: str):
    wb = Workbook()
    ws = wb.active
    ws.title = "Detail"

    # Styles
    hdr_font      = Font(name='Arial', bold=True, color='FFFFFF', size=11)
    hdr_fill      = PatternFill('solid', start_color='1F4E79')
    hdr_align     = Alignment(horizontal='center', vertical='center', wrap_text=True)
    tcode_font    = Font(name='Courier New', bold=True, color='1F4E79', size=10)
    others_font   = Font(name='Courier New', bold=True, color='C00000', size=10)
    doc_font      = Font(name='Arial', size=10)
    sentence_font = Font(name='Arial', size=10, italic=True)
    grey_font     = Font(name='Arial', size=10, italic=True, color='808080')
    top_align     = Alignment(horizontal='left', vertical='top',   wrap_text=True)
    mid_align     = Alignment(horizontal='center', vertical='center', wrap_text=True)

    def mk_border(weight='thin'):
        s = Side(style=weight, color='BDD7EE' if weight == 'thin' else '1F4E79')
        return Border(left=s, right=s, top=s, bottom=s)

    thin_border   = mk_border('thin')
    medium_border = mk_border('medium')
    alt_fill      = PatternFill('solid', start_color='EBF3FB')
    base_fill     = PatternFill('solid', start_color='FFFFFF')

    shared_styles = (hdr_font, hdr_fill, hdr_align, medium_border, thin_border)

    # Header row — Sheet 1
    for col, h in enumerate(['Transaction Code', 'Document Name', 'Source Sentence'], 1):
        c = ws.cell(row=1, column=col, value=h)
        c.font = hdr_font; c.fill = hdr_fill
        c.alignment = hdr_align; c.border = medium_border
    ws.row_dimensions[1].height = 32

    row   = 2
    block = 0

    for tcode in sorted(tcode_data.keys()):
        entries   = tcode_data[tcode]
        n         = len(entries)
        start_row = row
        fill      = alt_fill if (block % 2 == 0) else base_fill
        block    += 1

        for rel_path, sentence in entries:
            c = ws.cell(row=row, column=2, value=rel_path)
            c.font = doc_font; c.fill = fill
            c.alignment = top_align; c.border = thin_border

            c = ws.cell(row=row, column=3, value=sentence)
            c.font = sentence_font; c.fill = fill
            c.alignment = top_align; c.border = thin_border

            ws.row_dimensions[row].height = 38
            row += 1

        # Column A — write once, merge if multiple rows
        c = ws.cell(row=start_row, column=1, value=tcode)
        c.font = tcode_font; c.fill = fill
        c.alignment = mid_align; c.border = medium_border

        if n > 1:
            ws.merge_cells(start_row=start_row, start_column=1,
                           end_row=row - 1, end_column=1)
            c = ws.cell(row=start_row, column=1)
            c.font = tcode_font; c.fill = fill
            c.alignment = mid_align; c.border = medium_border

    # Others section
    if others:
        others_fill = PatternFill('solid', start_color='FFE7E7')
        start_row   = row
        n           = len(others)

        for rel_path in sorted(others):
            fill = alt_fill if (row % 2 == 0) else base_fill

            c = ws.cell(row=row, column=2, value=rel_path)
            c.font = doc_font; c.fill = fill
            c.alignment = top_align; c.border = thin_border

            c = ws.cell(row=row, column=3, value='No SAP transaction code detected')
            c.font = grey_font; c.fill = fill
            c.alignment = top_align; c.border = thin_border

            ws.row_dimensions[row].height = 30
            row += 1

        c = ws.cell(row=start_row, column=1, value='Others')
        c.font = others_font; c.fill = others_fill
        c.alignment = mid_align; c.border = medium_border

        if n > 1:
            ws.merge_cells(start_row=start_row, start_column=1,
                           end_row=row - 1, end_column=1)
            c = ws.cell(row=start_row, column=1)
            c.font = others_font; c.fill = others_fill
            c.alignment = mid_align; c.border = medium_border

    ws.column_dimensions['A'].width = 26
    ws.column_dimensions['B'].width = 38
    ws.column_dimensions['C'].width = 82
    ws.freeze_panes = 'B2'

    # Build Sheet 2 and Sheet 3
    build_sheet2(wb, tcode_data, others, shared_styles)
    build_sheet3(wb, tcode_data, others, shared_styles)
    build_sheet4(wb, tcode_data, others, None, shared_styles)

    wb.save(output_path)
    print(f"\nExcel saved to: {output_path}")
    print(f"  {len(tcode_data)} unique transaction code(s) found.")
    print(f"  {len(others)} document(s) with no codes -> 'Others'.")


def main():
    script_dir  = os.path.dirname(os.path.abspath(__file__))
    input_dir   = os.path.join(script_dir, 'input')
    output_dir  = os.path.join(script_dir, 'output')
    output_file = os.path.join(output_dir, 'SAP_Tcodes_Extract.xlsx')

    os.makedirs(input_dir,  exist_ok=True)
    os.makedirs(output_dir, exist_ok=True)

    print(f"Input  folder : {input_dir}")
    print(f"Output folder : {output_dir}\n")

    tcode_data, others = scan_folder(input_dir)
    build_excel(tcode_data, others, output_file)


if __name__ == '__main__':
    main()
