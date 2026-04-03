"""
Extract "Process Risk" tables from all .doc / .docx files in INPUT_FOLDER
(including sub-folders) and write a formatted Excel file to OUTPUT_FILE.

Layout  (mirrors the reference process_risk_output.xlsx format)
------
A  Folder Name        repeated on every row (no merging), grey fill
B  Word Doc Name      repeated on every row (no merging), grey fill
C+ Header Content N   per-file sub-header row (dark blue) then data rows (white)
                       N/A cells shaded for files with fewer columns

Notes
-----
- Supports both .docx and legacy .doc files.
  Legacy .doc: tried via LibreOffice first; falls back to pure-Python olefile.
- Merged title rows (e.g. "Risk and Issues Log") auto-detected and skipped.
- Rows where the Reference column (col 0) is blank are skipped.
- Files with no Process Risk section/table appear in a "Not Found" sheet.
- A "Summary" sheet shows distinct file counts from both sheets.

Usage
-----
1. Place .doc / .docx files in  <script_dir>/Input/  (sub-folders OK).
2. Run: python extract_process_risk_v2.py
3. Output: <script_dir>/Output/process_risk_output.xlsx
"""

import os
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

#  paths 
_HERE        = os.path.dirname(os.path.abspath(__file__))
INPUT_FOLDER = os.path.join(_HERE, "Input")
OUTPUT_FILE  = os.path.join(_HERE, "Output", "process_risk_output.xlsx")

#  constants 
RISK_KEYWORDS = ("process risk",)
HEADING_TAGS  = ("heading", "title", "toc")

W_T      = qn("w:t")
W_PPR    = qn("w:pPr")
W_PSTYLE = qn("w:pStyle")

#  styles 
def _F(h):          return PatternFill("solid", fgColor=h)
def _Ft(b, c):      return Font(name="Calibri", bold=b, color=c, size=11)
def _A(h, v):       return Alignment(horizontal=h, vertical=v, wrap_text=True)
def _S(s, c):       return Side(style=s, color=c)
def _B(t, b, l, r): return Border(top=t, bottom=b, left=l, right=r)

_NO, _MED, _THIN = Side(style=None), _S("medium", "4472C4"), _S("thin", "BFBFBF")
_TALL = _B(_THIN, _THIN, _THIN, _THIN)

F_HDR = _F("1F4E79")   # dark blue   global header row + per-file sub-header
F_FN  = _F("F2F2F2")   # light grey  folder/filename cells
F_VAL = _F("FFFFFF")   # white       data cells
F_NA  = _F("F5F5F5")   # near-white  N/A filler cells

FT_HDR = _Ft(True,  "FFFFFF")
FT_FN  = _Ft(True,  "1F4E79")
FT_VAL = _Ft(False, "000000")
FT_NA  = _Ft(False, "D0D0D0")

AL_CC = _A("center", "center")
AL_LC = _A("left",   "center")
AL_LT = _A("left",   "top")

#  cell writers 
#  cell value sanitiser 
import re as _re
_ILLEGAL_XML = _re.compile('[\x01-\x08\x0b\x0c\x0e-\x1f]')

def _clean(v):
    """Replace \x0b (Word soft-return) with newline; strip other XML-illegal chars."""
    if not isinstance(v, str):
        return v
    v = v.replace('\x0b', '\n')
    return _ILLEGAL_XML.sub('', v)

def _put(c, fill, font, align, border, value=None):
    c.fill = fill; c.font = font; c.alignment = align; c.border = border
    if value is not None:
        c.value = _clean(value)

def _hdr(c, v):   _put(c, F_HDR, FT_HDR, AL_CC, _TALL, v)   # dark blue header
def _sh(c, v):    _put(c, F_HDR, FT_HDR, AL_LC, _TALL, v)   # dark blue sub-header
def _fn(c, v):    _put(c, F_FN,  FT_FN,  AL_LC, _TALL, v)   # grey folder/filename
def _val(c, v=""): _put(c, F_VAL, FT_VAL, AL_LT, _TALL, v)  # white data
def _na(c):       _put(c, F_NA,  FT_NA,  AL_CC, _TALL, "")  # N/A filler

#  data model 
@dataclass
class Block:
    folder:       str
    filename:     str
    has_table:    bool = False
    bad_structure: bool = False
    base_hdrs:    list = field(default_factory=list)
    rows:         list = field(default_factory=list)
    col_list:     list = field(default_factory=list)
    col_widths:   list = field(default_factory=list)

#  docx helpers 
def _paras(cell) -> list:
    return [p.text.strip() for p in cell.paragraphs if p.text.strip()]

def _text(cell) -> str:
    return "\n".join(_paras(cell))

def _is_merged_title_row(row) -> bool:
    """Return True when this row is a title/spacer that should be skipped.

    Handles three cases:
    1. Fully merged row (all cells the same XML element) - whether empty or
       containing a title like "Risk and Issues Log"
    2. Multiple cells but all containing the same non-empty text
    3. Multiple cells that are all empty (blank spacer row)
    """
    dedup = _unique_cells(row)
    # Case 1: only one unique cell element (fully merged row) - always skip
    if len(dedup) == 1:
        return True
    # Case 2: all cells share the same non-empty text
    texts  = [_text(c).strip() for c in row.cells]
    unique = {t for t in texts if t}
    if len(unique) == 1:
        return True
    # Case 3: all cells are empty (blank spacer row)
    if not any(texts):
        return True
    return False

def _is_risk_section(body: list, tbl_elem) -> bool:
    """
    Scan upward through preceding paragraphs.
    Stop at the first heading-styled paragraph OR any paragraph whose text
    contains a risk keyword. Return True if that paragraph has a risk keyword.
    """
    try:
        idx = body.index(tbl_elem)
    except ValueError:
        return False
    for elem in reversed(body[:idx]):
        if elem.tag.split("}")[-1] != "p":
            continue
        pPr   = elem.find(W_PPR)
        style = ""
        if pPr is not None:
            ps = pPr.find(W_PSTYLE)
            if ps is not None:
                style = (ps.get(qn("w:val")) or "").lower().replace("-", " ")
        text       = "".join(t.text or "" for t in elem.iter() if t.tag == W_T).lower().strip()
        is_heading = any(h in style for h in HEADING_TAGS)
        has_risk   = any(k in text for k in RISK_KEYWORDS)
        if is_heading or (text and has_risk):
            return has_risk
    return False

def _unique_cells(row) -> list:
    """Return deduplicated cells — merged cells in Word repeat the same
    underlying XML element, causing duplicate column values. Only the
    first occurrence of each cell element is kept."""
    seen = set()
    result = []
    for cell in row.cells:
        if id(cell._tc) not in seen:
            seen.add(id(cell._tc))
            result.append(cell)
    return result


def _extract(table) -> tuple:
    """
    Parse table, skipping merged/blank title rows.
    Deduplicates merged cells so a column spanning 2 cells is counted once.

    Returns (base_hdrs, logical_rows).
    Returns (None, None) when the table structure is invalid (e.g. header row
    is fully merged into one cell, or first header cell has no real text).
    Rows where Reference (col 0) is blank or all columns are merged are skipped.
    """
    if not table.rows:
        return [], []

    rows = table.rows

    # Skip leading rows that are merged/blank title rows
    start = 0
    while start < len(rows) and _is_merged_title_row(rows[start]):
        start += 1

    if start >= len(rows):
        return [], []

    # Deduplicate header row cells
    hdr_cells = _unique_cells(rows[start])

    # If header row itself collapses to a single merged cell, table structure is bad
    if len(hdr_cells) <= 1:
        return None, None   # signals "Table structure not correct"

    base_hdrs = [
        _text(c).strip() or f"Col{i+1}"
        for i, c in enumerate(hdr_cells)
    ]

    # If first header is still a fallback name, table has no usable header text
    if base_hdrs[0].startswith("Col") and base_hdrs[0][3:].isdigit():
        return None, None   # signals "Table structure not correct"

    ref_hdr = base_hdrs[0]

    logical = []
    for row in rows[start + 1:]:
        # Skip rows where all columns are merged into one cell
        data_cells = _unique_cells(row)
        if len(data_cells) <= 1:
            continue

        rd = {
            h: _text(data_cells[i]).strip() if i < len(data_cells) else ""
            for i, h in enumerate(base_hdrs)
        }
        # Skip blank or "Reference Note" footer rows
        ref = rd.get(ref_hdr, "").strip()
        if not ref or ref.lower().startswith("reference note"):
            continue
        logical.append(rd)

    return base_hdrs, logical

#  Pure-Python OLE .doc extractor (fallback  no LibreOffice needed) 
def _extract_doc_ole(path: Path):
    """
    Extract Process Risk table directly from OLE binary .doc using olefile.
    Pure Python, no admin rights. pip install olefile
    Returns (headers, data_rows) or (None, None).
    """
    try:
        import olefile as _ole
    except ImportError:
        print("  [ERROR] olefile not installed. Run: pip install olefile")
        return None, None

    HEADING_STOP = re.compile(
        r"\r(POLICIES|PROCESS PERFORMANCE|PROCESS CONTACT|RECORD RETENTION|APPENDIX)",
        re.IGNORECASE,
    )
    try:
        ole = _ole.OleFileIO(str(path))
        raw = ole.openstream("WordDocument").read()
        ole.close()
    except Exception as exc:
        print(f"  [WARN] olefile could not open: {exc}")
        return None, None

    text = raw.decode("cp1252", errors="replace")

    heading_matches = list(re.finditer(r"PROCESS RISKS? AND MITIGATION", text, re.IGNORECASE))
    if not heading_matches:
        return None, None

    section_start = heading_matches[-1].end()
    stop_match    = HEADING_STOP.search(text, section_start)
    section_end   = stop_match.start() if stop_match else section_start + 3000
    section_text  = text[section_start:section_end]

    tbl_offset = section_text.find("\x07")
    if tbl_offset < 0:
        return None, None

    tbl_text  = section_text[tbl_offset:]
    raw_rows  = tbl_text.split("\x07\x07")
    rows_cells = []
    for raw_row in raw_rows:
        if not raw_row.strip("\r\n\x00 "):
            continue
        cells = [c.strip("\r\n\x00\x01\x13\x15 ") for c in raw_row.split("\x07")]
        cells = [c for c in cells if c.strip()]
        if cells:
            rows_cells.append(cells)

    if not rows_cells:
        return None, None

    def _is_merged(row):
        return len({c.strip() for c in row}) == 1 and len(row) >= 2

    start     = 1 if _is_merged(rows_cells[0]) else 0
    if start >= len(rows_cells):
        return None, None

    headers = rows_cells[start]
    n_cols  = len(headers)
    data_rows = []
    for row in rows_cells[start + 1:]:
        row = (row + [""] * n_cols)[:n_cols]
        rd  = dict(zip(headers, row))
        ref = row[0].strip() if row else ""
        if not ref or ref.lower().startswith("reference note"):
            continue
        data_rows.append(rd)

    return headers, data_rows

#  .doc  .docx conversion via LibreOffice 
def _soffice_exe() -> str:
    import shutil as _sh, sys as _sys
    for candidate in ("soffice", "soffice.exe"):
        if _sh.which(candidate):
            return candidate
    if _sys.platform == "win32":
        for base in (
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ):
            if Path(base).exists():
                return base
    return "soffice"

def _convert_doc_to_docx(src: Path):
    tmp_dir = tempfile.mkdtemp(prefix="doc_convert_")
    exe = _soffice_exe()
    try:
        result = subprocess.run(
            [exe, "--headless", "--convert-to", "docx", "--outdir", tmp_dir, str(src)],
            capture_output=True, text=True, timeout=120,
        )
        if result.returncode != 0:
            print(f"  [WARN] soffice failed (exe={exe}): {result.stderr.strip()}")
            shutil.rmtree(tmp_dir, ignore_errors=True)
            return None, None
        converted = next(Path(tmp_dir).glob("*.docx"), None)
        if converted is None:
            shutil.rmtree(tmp_dir, ignore_errors=True)
            return None, None
        return converted, tmp_dir
    except FileNotFoundError:
        print(f"  [INFO] LibreOffice not found  will use pure-Python OLE reader.")
        shutil.rmtree(tmp_dir, ignore_errors=True)
        return None, None
    except Exception as exc:
        print(f"  [WARN] conversion error: {exc}")
        shutil.rmtree(tmp_dir, ignore_errors=True)
        return None, None

def _open_doc(fpath: Path):
    if fpath.suffix.lower() == ".docx":
        return Document(str(fpath)), None
    print("   Converting .doc  .docx via LibreOffice ")
    converted, tmp_dir = _convert_doc_to_docx(fpath)
    if converted is not None:
        return Document(str(converted)), tmp_dir
    return None, None

#  block collection 
def collect() -> tuple:
    """Return (blocks, skipped, failed) where:
       blocks  = files successfully extracted
       skipped = (folder, filename, reason) for section/table not found
       failed  = (folder, filename) for files that could not be opened
    """
    blocks  = []
    skipped = []
    failed  = []   # files that could not be opened at all

    all_files = [
        p for p in sorted(Path(INPUT_FOLDER).rglob("*"))
        if p.suffix.lower() in (".doc", ".docx")
        and not p.name.startswith("~$")
        and p.is_file()
    ]
    stem_map: dict = {}
    for p in all_files:
        key = (str(p.parent), p.stem.lower())
        if key not in stem_map or p.suffix.lower() == ".docx":
            stem_map[key] = p
    found = sorted(stem_map.values(), key=lambda p: (str(p.parent), p.name))

    for fpath in found:
        rel    = fpath.relative_to(INPUT_FOLDER)
        folder = str(rel.parent) if str(rel.parent) != "." else "(root)"
        print(f"Processing: {rel}")

        blk           = Block(folder=folder, filename=fpath.name)
        doc, tmp_dir  = None, None
        found_section = False

        try:
            doc, tmp_dir = _open_doc(fpath)
        except Exception as exc:
            print(f"  [ERROR] open failed: {exc}")
            failed.append((folder, fpath.name))
            continue

        if doc is not None:
            #  python-docx path 
            try:
                body = list(doc.element.body)
                for ti, table in enumerate(doc.tables):
                    if not _is_risk_section(body, table._tbl):
                        continue
                    found_section = True
                    b_hdrs, l_rows = _extract(table)
                    if b_hdrs is None:
                        # Table structure is invalid (fully merged header, Col1 fallback)
                        print(f"   Table #{ti+1}: bad structure -- marking as error.")
                        blk.bad_structure = True
                        break
                    if not b_hdrs:
                        continue
                    blk.has_table  = True
                    blk.base_hdrs  = b_hdrs
                    blk.rows       = l_rows
                    blk.col_list   = b_hdrs
                    blk.col_widths = [min(max(len(l) * 1.2, 16), 55) for l in b_hdrs]
                    print(f"   Table #{ti+1}: {len(l_rows)} row(s), headers: {b_hdrs}")
                    break

                if not blk.has_table and not found_section:
                    for elem in body:
                        if elem.tag.split("}")[-1] != "p":
                            continue
                        pPr = elem.find(W_PPR); style = ""
                        if pPr is not None:
                            ps = pPr.find(W_PSTYLE)
                            if ps is not None:
                                style = (ps.get(qn("w:val")) or "").lower()
                        if not any(h in style for h in HEADING_TAGS):
                            continue
                        text_h = "".join(
                            t.text or "" for t in elem.iter() if t.tag == W_T
                        ).lower().strip()
                        if any(k in text_h for k in RISK_KEYWORDS):
                            found_section = True
                            break

            except Exception as exc:
                print(f"  [ERROR] parse failed: {exc}")
            finally:
                if tmp_dir:
                    shutil.rmtree(tmp_dir, ignore_errors=True)

        elif fpath.suffix.lower() == ".doc":
            #  pure-Python OLE fallback 
            print("   Trying pure-Python OLE reader ")
            try:
                b_hdrs, l_rows = _extract_doc_ole(fpath)
                if b_hdrs and l_rows:
                    found_section  = True
                    blk.has_table  = True
                    blk.base_hdrs  = b_hdrs
                    blk.rows       = l_rows
                    blk.col_list   = b_hdrs
                    blk.col_widths = [min(max(len(l) * 1.2, 16), 55) for l in b_hdrs]
                    print(f"   OLE: {len(l_rows)} row(s), headers: {b_hdrs}")
                elif b_hdrs is None:
                    found_section = False
                else:
                    found_section = True
            except Exception as exc:
                print(f"  [ERROR] OLE fallback failed: {exc}")

        # For .docx where doc stayed None (e.g. corrupt file not caught above)
        if doc is None and fpath.suffix.lower() == ".docx":
            if (folder, fpath.name) not in failed:
                failed.append((folder, fpath.name))
            continue

        if blk.bad_structure:
            reason = "Table Structure Not Correct"
            print(f"  [INFO] {reason} -- adding to skipped list.")
            skipped.append((folder, fpath.name, reason))
            continue

        if not blk.has_table:
            if found_section:
                reason = "Table Not Found"
            else:
                reason = "Process Risk Section Not Found"
            print(f"  [INFO] {reason} -- adding to skipped list.")
            skipped.append((folder, fpath.name, reason))
            continue

        blocks.append(blk)

    return blocks, skipped, failed

#  excel writer 
CS = 3   # content starts at column C (1=A, 2=B, 3=C )

def write(blocks: list, path: str, skipped: list | None = None,
          failed: list | None = None) -> None:
    skipped = skipped or []
    failed  = failed  or []
    if not blocks and not skipped and not failed:
        print("[WARN] Nothing to write.")
        return

    # Union of all column headers across files (in order of first appearance)
    all_hdr_sets = [tuple(b.col_list) for b in blocks if b.has_table]
    if all_hdr_sets:
        seen: dict = {}
        for hset in all_hdr_sets:
            for h in hset:
                if h not in seen:
                    seen[h] = True
        common_hdrs = list(seen.keys())
    else:
        common_hdrs = ["Reference", "Risks  What Could Go Wrong?",
                       "Implication", "Control No. per RACF"]

    max_cc     = max(len(b.col_list) for b in blocks) if blocks else len(common_hdrs)
    total_cols = CS - 1 + max_cc

    wb = Workbook()
    ws = wb.active
    ws.title = "Process Risk"

    #  Row 1: global header 
    _hdr(ws.cell(1, 1), "Folder Name")
    _hdr(ws.cell(1, 2), "Word Doc Name")
    for i in range(max_cc):
        _hdr(ws.cell(1, CS + i), f"Header Content {i+1}")

    cur = 2
    for blk in blocks:
        n_own = len(blk.col_list)

        #  Per-file sub-header row + data rows 
        # The sub-header row (col_list labels like "Reference", "Risks...", etc.)
        # is written inline with data rows below.
        # Rule: if Header Content 1 value == "Reference" (or the actual ref header
        # name from this file) treat the whole row as a sub-header (dark blue);
        # all other rows get white background.
        ref_col_name = blk.base_hdrs[0] if blk.base_hdrs else ""

        # Sub-header row (column name labels)
        _fn(ws.cell(cur, 1), blk.folder)
        _fn(ws.cell(cur, 2), blk.filename)
        for i, lbl in enumerate(blk.col_list):
            _sh(ws.cell(cur, CS + i), lbl)
        for ci in range(CS + n_own, total_cols + 1):
            _sh(ws.cell(cur, ci), "")
        cur += 1

        # Data rows: white background throughout
        for lr in blk.rows:
            _fn(ws.cell(cur, 1), blk.folder)
            _fn(ws.cell(cur, 2), blk.filename)
            ref_val = lr.get(ref_col_name, "").strip()
            # If the reference value itself equals the header name (e.g. "Reference")
            # treat this row as a sub-header row with dark-blue background
            if ref_val.lower() == ref_col_name.lower():
                for i, h in enumerate(blk.base_hdrs):
                    _sh(ws.cell(cur, CS + i), lr.get(h, ""))
                for ci in range(CS + n_own, total_cols + 1):
                    _sh(ws.cell(cur, ci), "")
            else:
                # Normal data row — fully white
                for i, h in enumerate(blk.base_hdrs):
                    _val(ws.cell(cur, CS + i), lr.get(h, ""))
                for ci in range(CS + n_own, total_cols + 1):
                    _na(ws.cell(cur, ci))
            cur += 1

    ws.auto_filter.ref = f"A1:{get_column_letter(total_cols)}{cur - 1}"
    ws.freeze_panes    = "A2"

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 48
    widths: dict = {}
    for blk in blocks:
        for i, w in enumerate(blk.col_widths):
            widths[i] = max(widths.get(i, 0), w)
    for i, w in widths.items():
        ws.column_dimensions[get_column_letter(CS + i)].width = w

    ws.row_dimensions[1].height = 28
    for r in range(2, cur):
        ws.row_dimensions[r].height = 36

    # ── Shared styles for error sheets ──────────────────────────────────────
    thin     = Side(style="thin",   color="BFBFBF")
    bdr_cell = Border(top=thin, bottom=thin, left=thin, right=thin)
    al_hdr   = _A("center", "center")
    al_val   = _A("left",   "center")

    def _err_sheet(wb, title, col_headers, rows, hdr_color,
                   row_colors=None):
        """Generic helper to build a 3-col error/info sheet."""
        ws = wb.create_sheet(title)
        h_fill = _F(hdr_color); h_font = _Ft(True, "FFFFFF")
        v_fill = _F("FFFFFF");  v_font = _Ft(False, "000000")

        for ci, lbl in enumerate(col_headers, 1):
            c = ws.cell(1, ci)
            c.value = lbl; c.fill = h_fill; c.font = h_font
            c.alignment = al_hdr; c.border = bdr_cell
        ws.row_dimensions[1].height = 28

        for ri, row_vals in enumerate(rows, 2):
            rc = (row_colors[ri-2] if row_colors else None) or "FFFFFF"
            rf = _F(rc)
            for ci, val in enumerate(row_vals, 1):
                c = ws.cell(ri, ci)
                c.value = val; c.border = bdr_cell; c.alignment = al_val
                c.fill = rf
                c.font = _Ft(True, "7F6000") if (rc != "FFFFFF" and ci == len(row_vals)) else v_font
            ws.row_dimensions[ri].height = 36

        ws.column_dimensions["A"].width = 22
        ws.column_dimensions["B"].width = 50
        if len(col_headers) >= 3:
            ws.column_dimensions["C"].width = 45
        if rows:
            ws.auto_filter.ref = f"A1:{get_column_letter(len(col_headers))}{1 + len(rows)}"
        ws.freeze_panes = "A2"
        return ws

    # ── Sheet 2: Section-Table Not Found ─────────────────────────────────────
    if skipped:
        def _row_color(reason):
            if reason == "Table Not Found":
                return "FFF2CC"         # amber
            if reason == "Table Structure Not Correct":
                return "FCE4D6"         # light orange
            return "FFFFFF"             # white

        row_colors = [_row_color(r[2]) for r in skipped]
        _err_sheet(
            wb,
            title="Section-Table Not Found",
            col_headers=["Folder Name", "File Name", "Error"],
            rows=skipped,
            hdr_color="C00000",
            row_colors=row_colors,
        )
        print(f"   Section-Table Not Found sheet: {len(skipped)} file(s) listed")

    # ── Sheet 3: File Open Error ──────────────────────────────────────────────
    if failed:
        _err_sheet(
            wb,
            title="File Open Error",
            col_headers=["Folder Name", "File Name"],
            rows=failed,
            hdr_color="833C00",
        )
        print(f"   File Open Error sheet: {len(failed)} file(s) listed")

    # ── Sheet 4: Summary ──────────────────────────────────────────────────────
    ws_sum = wb.create_sheet("Summary")
    distinct_s1 = len({b.filename for b in blocks})
    distinct_s2 = len({r[1] for r in skipped}) if skipped else 0
    distinct_s3 = len({r[1] for r in failed})  if failed  else 0

    s_hdr_fill   = _F("1F4E79"); s_hdr_font   = _Ft(True,  "FFFFFF")
    s_lbl_fill   = _F("2E75B6"); s_lbl_font   = _Ft(True,  "FFFFFF")
    s_val_fill   = _F("FFFFFF"); s_val_font   = _Ft(True,  "1F4E79")
    s_tot_fill   = _F("D6E4F0"); s_tot_font   = _Ft(True,  "1F4E79")
    s_thin = Side(style="thin",   color="BFBFBF")
    s_med  = Side(style="medium", color="4472C4")
    bdr_h  = Border(top=s_med,  bottom=s_med,  left=s_med,  right=s_med)
    bdr_l  = Border(top=s_thin, bottom=s_thin, left=s_med,  right=s_thin)
    bdr_v  = Border(top=s_thin, bottom=s_thin, left=s_thin, right=s_med)
    bdr_t  = Border(top=s_med,  bottom=s_med,  left=s_med,  right=s_med)
    al_cc  = _A("center", "center")
    al_lc  = _A("left",   "center")

    def _s_hdr(c, v): c.value=v; c.fill=s_hdr_fill; c.font=s_hdr_font; c.alignment=al_cc; c.border=bdr_h
    def _s_lbl(c, v): c.value=v; c.fill=s_lbl_fill; c.font=s_lbl_font; c.alignment=al_lc; c.border=bdr_l
    def _s_val(c, v): c.value=v; c.fill=s_val_fill; c.font=s_val_font; c.alignment=al_cc; c.border=bdr_v

    _s_hdr(ws_sum.cell(1,1), "Sheet Name")
    _s_hdr(ws_sum.cell(1,2), "Description")
    _s_hdr(ws_sum.cell(1,3), "Distinct File Count")

    all_summary_rows = [
        ("Process Risk",            "Files with Process Risk table extracted",  distinct_s1),
        ("Section-Table Not Found", "Files where section or table not found",   distinct_s2),
        ("File Open Error",         "Files that could not be opened",           distinct_s3),
    ]
    # Only include rows where Distinct File Count > 0
    summary_rows = [(s, d, c) for s, d, c in all_summary_rows if c > 0]

    for ri, (sheet, desc, cnt) in enumerate(summary_rows, 2):
        _s_lbl(ws_sum.cell(ri, 1), sheet)
        _s_lbl(ws_sum.cell(ri, 2), desc)
        _s_val(ws_sum.cell(ri, 3), cnt)

    total_row = 2 + len(summary_rows)
    total_cnt = sum(c for _, _, c in summary_rows)
    for ci, val in enumerate(["Total", "All files processed", total_cnt], 1):
        c = ws_sum.cell(total_row, ci)
        c.value = val; c.fill = s_tot_fill; c.font = s_tot_font
        c.border = bdr_t
        c.alignment = al_lc if ci < 3 else al_cc

    ws_sum.column_dimensions["A"].width = 28
    ws_sum.column_dimensions["B"].width = 45
    ws_sum.column_dimensions["C"].width = 22
    for r in range(1, total_row + 1):
        ws_sum.row_dimensions[r].height = 32

    os.makedirs(os.path.dirname(path), exist_ok=True)
    wb.save(path)
    print(f"\n  {cur - 1} rows | {total_cols} cols  {path}")


#  entry point 
if __name__ == "__main__":
    blocks, skipped, failed = collect()
    write(blocks, OUTPUT_FILE, skipped, failed)