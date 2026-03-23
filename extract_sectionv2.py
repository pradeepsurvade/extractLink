"""
extract_section16.py
Scans .docx files in 'input/' (recursive), extracts Section 16 tables,
writes formatted results to 'output/Section16_Extract.xlsx'.

Sheet 1 "All Tables"    : consolidated view — all docs, all tables
Sheet 2+ (per table)    : one sheet per Word table found across all docs,
                          named by table header row; each sheet has
                          Folder Name + Word Doc Name prepended as col A & B.

Requirements: pip install python-docx openpyxl
"""

import sys
from pathlib import Path
from dataclasses import dataclass, field

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table as DocxTable
except ImportError:
    sys.exit("Missing dependency: pip install python-docx")

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
except ImportError:
    sys.exit("Missing dependency: pip install openpyxl")

INPUT_DIR  = Path(__file__).parent / "input"
OUTPUT_DIR = Path(__file__).parent / "output"

# ── Style constants ───────────────────────────────────────────────────────────
_THIN   = Side(style="thin",   color="BFBFBF")
_MEDIUM = Side(style="medium", color="4472C4")
_BRD         = Border(left=_THIN,   right=_THIN,   top=_THIN,   bottom=_THIN)
_BRD_META    = Border(left=_MEDIUM, right=_MEDIUM, top=_MEDIUM, bottom=_MEDIUM)
_BRD_DIVIDER = Border(left=_THIN,   right=_THIN,   bottom=_THIN,
                      top=Side(style="medium", color="1F4E79"))

_F_SHEET_HDR = Font(name="Arial", bold=True, color="FFFFFF", size=11)
_F_TBL_HDR   = Font(name="Arial", bold=True, color="FFFFFF", size=10)
_F_DATA      = Font(name="Arial", size=10)
_F_META      = Font(name="Arial", bold=True, color="1F4E79", size=10)

_FILL_SHEET_HDR = PatternFill("solid", fgColor="1F4E79")
_FILL_TBL_HDR   = PatternFill("solid", fgColor="2E75B6")
_FILL_ALT       = PatternFill("solid", fgColor="D9E8F5")
_FILL_META      = PatternFill("solid", fgColor="F2F2F2")
_FILL_NONE      = PatternFill()

_ALIGN_CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)
_ALIGN_TOP    = Alignment(horizontal="left",   vertical="top",    wrap_text=True)
_ALIGN_MID    = Alignment(horizontal="left",   vertical="center", wrap_text=True)

def _style(cell, font=None, fill=None, align=None, border=None):
    if font:   cell.font      = font
    if fill:   cell.fill      = fill
    if align:  cell.alignment = align
    if border: cell.border    = border

# ── Data model ────────────────────────────────────────────────────────────────
@dataclass
class TableBlock:
    """A single table extracted from a Word doc."""
    folder:   str
    filename: str
    section:  str
    rows:     list[list[str]]   # first row = header, rest = data

@dataclass
class DocRecord:
    folder:   str
    filename: str
    section:  str
    tables:   list[TableBlock] = field(default_factory=list)

# ── Extraction ────────────────────────────────────────────────────────────────
def _pstyle(elem) -> str:
    ps = elem.find(".//" + qn("w:pStyle"))
    return ps.get(qn("w:val")) if ps is not None else ""

def _text(elem) -> str:
    return "".join(t.text or "" for t in elem.findall(".//" + qn("w:t"))).strip()

def _parse_table(tbl_elem, doc) -> list[list[str]]:
    rows = []
    for row in DocxTable(tbl_elem, doc).rows:
        seen, cells = set(), []
        for cell in row.cells:
            if id(cell) not in seen:
                seen.add(id(cell))
                cells.append(cell.text.replace("\n", " ").strip())
        rows.append(cells)
    return rows

def extract_section16(doc_path: Path) -> tuple[str, list[list[list[str]]]]:
    """Return (section_name, list_of_tables) where each table is a list of row-lists."""
    try:
        doc = Document(str(doc_path))
    except Exception as exc:
        return f"ERROR: {exc}", []

    section, tables, in_s16, h1_count = "", [], False, 0
    for elem in doc.element.body:
        tag = elem.tag.split("}")[-1]
        if tag == "p" and _pstyle(elem) == "Heading1":
            h1_count += 1
            if h1_count == 16:
                in_s16, section = True, "16\t" + _text(elem)
            elif h1_count == 17 and in_s16:
                break
        elif tag == "tbl" and in_s16:
            rows = _parse_table(elem, doc)
            if rows:
                tables.append(rows)

    return (section or "Section 16 Not Found"), tables

# ── Excel helpers ─────────────────────────────────────────────────────────────
def _sheet_header(ws, labels, row_height=28):
    for col, label in enumerate(labels, start=1):
        c = ws.cell(row=1, column=col, value=label)
        _style(c, font=_F_SHEET_HDR, fill=_FILL_SHEET_HDR, align=_ALIGN_CENTER, border=_BRD)
    ws.row_dimensions[1].height = row_height

def _set_col_widths(ws, widths: dict, default_start: int, default_width=32):
    for col_letter, w in widths.items():
        ws.column_dimensions[col_letter].width = w
    for col in range(default_start, ws.max_column + 1):
        if get_column_letter(col) not in widths:
            ws.column_dimensions[get_column_letter(col)].width = default_width

def _write_table_rows(ws, tbl_rows: list[list[str]], start_row: int,
                      meta_col_count: int = 0) -> int:
    """
    Write table rows into ws starting at start_row.
    meta_col_count: number of leading columns already written (skip fill for those).
    Returns the next available excel row.
    """
    excel_row = start_row
    alt = 0
    for i, trow in enumerate(tbl_rows):
        is_hdr   = (i == 0)
        fill_row  = _FILL_TBL_HDR if is_hdr else (_FILL_ALT if alt % 2 else _FILL_NONE)
        font_row  = _F_TBL_HDR    if is_hdr else _F_DATA
        align_row = _ALIGN_MID    if is_hdr else _ALIGN_TOP

        for offset, text in enumerate(trow):
            c = ws.cell(row=excel_row, column=meta_col_count + 1 + offset, value=text)
            _style(c, font=font_row, fill=fill_row, align=align_row, border=_BRD)

        # Fill meta cols with matching row colour
        for col in range(1, meta_col_count + 1):
            c = ws.cell(row=excel_row, column=col)
            c.fill, c.border = fill_row, _BRD

        if not is_hdr:
            alt += 1
        ws.row_dimensions[excel_row].height = 18
        excel_row += 1
    return excel_row

def _merge_col(ws, col, start_row, end_row, value, style_fn):
    if end_row > start_row:
        ws.merge_cells(start_row=start_row, start_column=col,
                       end_row=end_row,     end_column=col)
    c = ws.cell(row=start_row, column=col, value=value)
    style_fn(c)

# ── Sheet 1: consolidated view ────────────────────────────────────────────────
META_COLS = 3   # Folder | Doc Name | Section Name

def _write_consolidated(wb, records: list[DocRecord]):
    ws = wb.active
    ws.title = "All Tables"

    _sheet_header(ws, ["Folder Name", "Word Doc Name", "Section Name", "Table Content"])

    excel_row = 2
    for rec in records:
        meta_vals = [rec.folder, rec.filename, rec.section]

        if not rec.tables:
            for col, val in enumerate(meta_vals, start=1):
                c = ws.cell(row=excel_row, column=col, value=val)
                _style(c, font=_F_META, fill=_FILL_META, align=_ALIGN_TOP, border=_BRD_META)
            c = ws.cell(row=excel_row, column=META_COLS + 1, value="(No tables found in Section 16)")
            _style(c, font=_F_DATA, align=_ALIGN_TOP, border=_BRD)
            excel_row += 1
            continue

        doc_start = excel_row
        need_divider = False

        for tbl in rec.tables:
            if need_divider:
                # Apply thick-top border on the first row of the next table
                for col in range(1, META_COLS + len(tbl[0]) + 2):
                    c = ws.cell(row=excel_row, column=col)
                    c.border = _BRD_DIVIDER

            next_row = _write_table_rows(ws, tbl, excel_row, meta_col_count=META_COLS)
            excel_row    = next_row
            need_divider = True

        # Merge meta cols across all rows for this doc
        for col, val in enumerate(meta_vals, start=1):
            _merge_col(ws, col, doc_start, excel_row - 1, val,
                       lambda c: _style(c, font=_F_META, fill=_FILL_META,
                                        align=_ALIGN_TOP, border=_BRD_META))

    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 42
    ws.column_dimensions["C"].width = 38
    _set_col_widths(ws, {"A": 30, "B": 42, "C": 38}, default_start=4)
    ws.freeze_panes = "D2"

# ── Sheet 2+: one sheet per unique table (by header row) ─────────────────────
def _safe_sheet_name(name: str, existing: set[str]) -> str:
    """Strip invalid chars, truncate to 31 chars, ensure unique."""
    import re
    name = re.sub(r"[\/*?:\[\]]", "-", name)
    base = name[:28].strip()
    candidate, i = base, 1
    while candidate in existing:
        candidate = f"{base[:25]}_{i}"
        i += 1
    existing.add(candidate)
    return candidate

def _write_per_table_sheets(wb, records: list[DocRecord]):
    """
    Group all TableBlocks by their header row (first row of each table).
    Each unique header gets its own sheet.
    """
    # Collect all table blocks across all docs
    all_blocks: list[TableBlock] = []
    for rec in records:
        for tbl in rec.tables:
            all_blocks.append(TableBlock(
                folder=rec.folder, filename=rec.filename,
                section=rec.section, rows=tbl
            ))

    if not all_blocks:
        return

    # Group by header (first row joined as string key)
    from collections import defaultdict
    groups: dict[str, list[TableBlock]] = defaultdict(list)
    for blk in all_blocks:
        key = " | ".join(blk.rows[0]) if blk.rows else "Unknown"
        groups[key].append(blk)

    existing_names = {s.title for s in wb.worksheets}

    for header_key, blocks in groups.items():
        sheet_name = _safe_sheet_name(header_key, existing_names)
        ws = wb.create_sheet(title=sheet_name)

        # Determine all column headers: Folder Name, Word Doc Name, then table columns
        tbl_headers = blocks[0].rows[0] if blocks[0].rows else []
        all_headers = ["Folder Name", "Word Doc Name"] + tbl_headers
        _sheet_header(ws, all_headers)

        excel_row = 2
        alt = 0
        for blk in blocks:
            data_rows = blk.rows[1:]  # skip the table's own header row
            if not data_rows:
                continue
            for trow in data_rows:
                fill_row  = _FILL_ALT if alt % 2 else _FILL_NONE
                font_row  = _F_DATA
                align_row = _ALIGN_TOP

                # Col A: Folder, Col B: Doc name
                for col, val in enumerate([blk.folder, blk.filename], start=1):
                    c = ws.cell(row=excel_row, column=col, value=val)
                    _style(c, font=_F_META, fill=_FILL_META, align=_ALIGN_TOP, border=_BRD_META)

                # Col C onwards: table data
                for offset, text in enumerate(trow):
                    c = ws.cell(row=excel_row, column=3 + offset, value=text)
                    _style(c, font=font_row, fill=fill_row, align=align_row, border=_BRD)

                alt += 1
                ws.row_dimensions[excel_row].height = 18
                excel_row += 1

        _set_col_widths(ws, {"A": 30, "B": 42}, default_start=3)
        ws.freeze_panes = "C2"

# ── Main writer ───────────────────────────────────────────────────────────────
def write_excel(records: list[DocRecord], output_path: Path) -> None:
    wb = Workbook()
    _write_consolidated(wb, records)
    _write_per_table_sheets(wb, records)
    wb.save(str(output_path))

# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    if not INPUT_DIR.exists():
        sys.exit(f"Input folder not found: {INPUT_DIR}")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    docx_files = sorted(p for p in INPUT_DIR.rglob("*.docx") if not p.name.startswith("~$"))
    if not docx_files:
        sys.exit(f"No .docx files found under: {INPUT_DIR}")

    print(f"Found {len(docx_files)} Word file(s). Extracting Section 16 tables …\n")

    records = []
    for doc_path in docx_files:
        rel    = doc_path.relative_to(INPUT_DIR)
        folder = "" if str(rel.parent) == "." else str(rel.parent)
        section, tables = extract_section16(doc_path)
        n = sum(len(t) for t in tables)
        print(f"  {'✓' if tables else '–'}  {rel}  →  {section[:60]}  ({len(tables)} tables, {n} rows)")
        records.append(DocRecord(folder=folder, filename=doc_path.name,
                                 section=section, tables=tables))

    output_file = OUTPUT_DIR / "Section16_Extract.xlsx"
    write_excel(records, output_file)
    print(f"\nDone! Output saved to: {output_file}")

if __name__ == "__main__":
    main()