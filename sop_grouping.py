"""
SOP Document Grouping — Semantic Similarity Analysis
======================================================
Scans .docx files in input/, extracts target section content,
and groups documents by semantic similarity of that content.

  Column A: Group Name      (derived from process reasoning)
  Column B: Document Name
  Column C: Reason for Grouping
  Column D: SAP Transaction Codes found in the document

Usage:
  python sop_grouping.py               (auto-detects number of groups)
  python sop_grouping.py --groups 5    (force 5 groups)
"""

import os
import re
import math
import argparse
from collections import defaultdict

try:
    from docx import Document
except ImportError:
    raise SystemExit("python-docx not installed. Run: pip install python-docx")

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.cluster import AgglomerativeClustering
    from sklearn.metrics.pairwise import cosine_similarity as sklearn_cosine
    import numpy as np
except ImportError:
    raise SystemExit("scikit-learn / numpy not installed. Run: pip install scikit-learn numpy")

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
except ImportError:
    raise SystemExit("openpyxl not installed. Run: pip install openpyxl")


# ===========================================================================
# PROMPT — Configure everything about how grouping works
# ===========================================================================
#
# Write plain English. The prompt controls:
#   1. WHAT to compare      — which content to focus on
#   2. HOW to group         — what makes documents similar
#   3. GROUP BOUNDARIES     — what should or should not go together
#
# ── WHAT TO COMPARE (choose one, or leave blank for default) ────────────────
#   "compare only SAP transaction codes"
#   "compare only process steps"
#   "compare only activities overview"
#   "compare SAP codes and process steps together"
#   "compare full section content"           ← default (all TARGET_SECTIONS)
#
# ── HOW TO GROUP (describe the similarity criteria) ─────────────────────────
#   "group by the type of financial transaction being processed"
#   "group by the SAP module used (FI, CO, MM, HR, PS)"
#   "group documents that use the same approval workflow"
#   "group by the department responsible for the process"
#   "group by whether documents involve month-end closing activities"
#
# ── GROUP BOUNDARIES (explicit include/exclude rules) ───────────────────────
#   "GL reconciliation and journal posting should be in the same group"
#   "payroll and employee onboarding belong together as HR processes"
#   "keep asset depreciation and capitalisation in the same group"
#   "separate finance processes from HR and project management"
#
# You can combine all three types in one prompt. Examples:
#
# EXAMPLE 1 — Group by SAP module:
#   PROMPT = """
#   Compare only SAP transaction codes.
#   Group documents by the SAP module they use.
#   FI documents (GL, AP, AR, asset accounting) should be one group.
#   HR documents (payroll, onboarding) should be a separate group.
#   PS documents (project system, WBS, CJ codes) should be their own group.
#   """
#
# EXAMPLE 2 — Group by process flow similarity:
#   PROMPT = """
#   Compare SAP codes and process steps together.
#   Group documents that follow the same end-to-end process flow.
#   Documents with similar approval steps and posting sequences belong together.
#   Keep month-end closing processes separate from day-to-day transactions.
#   """
#
# EXAMPLE 3 — Group by department:
#   PROMPT = """
#   Group by the department responsible for the process.
#   Finance processes (invoicing, GL, assets) in one group.
#   HR processes (payroll, onboarding, benefits) in one group.
#   Project management processes (WBS, settlement, budgeting) in one group.
#   """
# ===========================================================================
PROMPT = """
Focus on process steps to find similarity.
Compare the detailed process steps across documents.
Group documents that follow the same process flow and step sequence.
Documents from different regions but with the same process steps belong in the same group.
Ignore region-specific terms, country names, and local variations — focus only on what the process does.
"""


# ---------------------------------------------------------------------------
# TARGET SECTIONS
# Headings to extract from each document for comparison.
# Add or remove section names (case-insensitive partial match).
# ---------------------------------------------------------------------------
TARGET_SECTIONS = [
    "SOP ACTIVITIES OVERVIEW",
    "DETAILED PROCESS STEPS",
    # "SCOPE",
    # "OBJECTIVES",
    # "ROLES AND RESPONSIBILITIES",
]


# ---------------------------------------------------------------------------
# EMBEDDING MODEL — runs 100% OFFLINE after one-time setup
#
# YOUR DATA NEVER GOES TO THE INTERNET:
#   - The model is downloaded ONCE to your local machine cache
#   - All processing happens entirely on your CPU, in local memory
#   - TRANSFORMERS_OFFLINE=1 is set in code — NO network calls during inference
#   - local_files_only=True — raises error instead of silently downloading
#
# ONE-TIME SETUP (internet needed only for this download step):
#   Step 1:  pip install sentence-transformers
#   Step 2:  python -c "from sentence_transformers import SentenceTransformer; SentenceTransformer('all-MiniLM-L6-v2')"
#   After that: run sop_grouping.py fully offline, forever.
#
# Model options (all open-source, Apache 2.0 / MIT licensed):
#   "all-MiniLM-L6-v2"          — recommended: best accuracy + speed (~90MB)
#   "all-MiniLM-L12-v2"         — slightly more accurate (~120MB)
#   "paraphrase-MiniLM-L3-v2"   — fastest, smallest (~60MB)
# ---------------------------------------------------------------------------
EMBEDDING_MODEL = "all-MiniLM-L6-v2"


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
MIN_DOCS            = 2
GROUP_NAME_KEYWORDS = 4
DOC_REASON_KEYWORDS = 6

# Minimum number of words a document's process content must have to be
# treated as a full SOP and included in process-similarity clustering.
# Documents below this threshold are labelled "Reference Documents"
# (quick references, checklists, summaries) and shown separately.
# Increase this value to be stricter about what counts as a full SOP.
MIN_PROCESS_WORDS   = 20

STOPWORDS = {
    'the','a','an','and','or','but','in','on','at','to','for','of','with',
    'by','from','is','are','was','were','be','been','being','have','has',
    'had','do','does','did','will','would','could','should','may','might',
    'shall','can','this','that','these','those','it','its','as','if','then',
    'than','so','not','no','nor','up','out','into','through','during',
    'before','after','above','below','between','each','all','both','few',
    'more','most','other','some','such','any','per','also','etc','please',
    'must','user','step','steps','click','select','enter','open','close',
    'following','using','used','use','new','save','screen','button','field',
    'fields','form','page','section','document','documents','file','files',
    'system','sap','transaction','tcode','code','process','procedure','go',
    'see','note','menu','via','within','without','where','when','how',
    'what','which',
}


# ---------------------------------------------------------------------------
# SAP T-code extraction (inlined)
# ---------------------------------------------------------------------------
_TCODE_TRIGGER_PATTERNS = [
    re.compile(r'Run\s+SAP\s+[Tt][-\s]?[Cc]ode\s+([A-Z][A-Z0-9_\-]{1,29})',    re.IGNORECASE),
    re.compile(r'Run\s+the\s+([A-Z][A-Z0-9_\-]{1,29})\s+[Tt]code',              re.IGNORECASE),
    re.compile(r'Run\s+([A-Z][A-Z0-9_\-]{1,29})\s+SAP\s+[Tt]code',              re.IGNORECASE),
    re.compile(r'Run\s+([A-Z][A-Z0-9_\-]{1,29})\s+(?:SAP|[Tt][-\s]?[Cc]ode|[Tt]ransaction)\b', re.IGNORECASE),
    re.compile(r'[Tt][-\s]?[Cc]ode[s]?\s*[:\(]?\s*([A-Z][A-Z0-9_\-]{1,29})\)?',re.IGNORECASE),
    re.compile(r'SAP\s+[Tt]ransaction\s*[:\(]?\s*([A-Z][A-Z0-9_\-]{1,29})\)?',  re.IGNORECASE),
    re.compile(r'SAP\s+\w[\w\s]{0,40}[Tt]ransaction\s*\(\s*([A-Z][A-Z0-9_\-]{1,29})\s*\)', re.IGNORECASE),
    re.compile(r'[Tt]ransaction\s+[Cc]ode\s*[:\(]?\s*([A-Z][A-Z0-9_\-]{1,29})\)?', re.IGNORECASE),
    re.compile(r'\bOpen\s+Transaction\s*[\u201c\u201d\u2018\u2019"\']?([A-Z][A-Z0-9_\-]{1,29})[\u201c\u201d\u2018\u2019"\']?', re.IGNORECASE),
    re.compile(r'\bOpen\s+Transaction\s+([A-Z][A-Z0-9_\-]{1,29})\b',            re.IGNORECASE),
    re.compile(r'\bTransaction\s+([A-Z][A-Z0-9_\-]{1,29})\b'),
    re.compile(r'\b(S_ALR_[A-Z0-9_]+)\b',                                        re.IGNORECASE),
]

_NOT_A_TCODE = {
    'THE','AND','FOR','RUN','SAP','CODE','CODES','WITH','FROM','INTO','THIS',
    'THAT','HAVE','BEEN','WILL','ALSO','CAN','NOT','ARE','ALL','BUT','USE',
    'NEW','OLD','END','ADD','SET','GET','PUT','OUT','OFF','TOP','BOX','YES',
    'NO','OK','GO','DO','IF','AS','AT','BE','BY','IN','IS','IT','OF','ON',
    'OR','SO','TO','UP','US','WE','ME','TCODE','TCODES','TRANSACTION',
    'REPORT','MODULE','SYSTEM','TABLE','FIELD','VALUE','PLEASE','CLICK',
    'OPEN','CLOSE','ENTER','SELECT','PRESS','BUTTON','SCREEN','WINDOW',
    'MENU','LIST','VIEW','NEXT','BACK','SAVE','EXIT','NOTE','STEP','THEN',
    'WHEN','ONCE','AFTER','BEFORE','USING','BELOW','ABOVE','RIGHT','LEFT',
    'HERE','THERE','DIFF','DEP',
}

_TABLE_CODE_HEADERS = {
    'process terms','process acronyms','process term','process acronym',
    'transaction code','transaction codes','t-code','tcode','t code',
    'sap transaction','sap tcode','sap t-code','sap code','sap codes',
}

_TABLE_IGNORE = {
    'NA','N/A','N.A','N.A.','NONE','NIL','TBD','TBC','-','--','---',
    'WBS','WNS','N','A','YES','NO','TRUE','FALSE',
}


def _is_valid_tcode(c: str) -> bool:
    c = c.upper().strip().rstrip('.,;)(')
    if not c or c in _NOT_A_TCODE or len(c) < 2 or not c[0].isalpha():
        return False
    if not re.match(r'^[A-Z][A-Z0-9_]+$', c):
        return False
    return True


def _strip_prefix(code: str) -> str:
    return code[7:] if code.upper().startswith('T_CODE_') else code


def _extract_tcodes_from_text(text: str) -> set:
    tcodes, seen = set(), set()
    for pat in _TCODE_TRIGGER_PATTERNS:
        for m in pat.finditer(text):
            raw = m.group(1).strip().rstrip('.,;)(').upper().replace('-', '_')
            raw = _strip_prefix(raw)
            if _is_valid_tcode(raw) and raw not in seen:
                seen.add(raw); tcodes.add(raw)
    return tcodes


def extract_sap_codes(filepath: str) -> list:
    doc = Document(filepath)
    tcodes = set()
    for para in doc.paragraphs:
        tcodes |= _extract_tcodes_from_text(para.text)
    for table in doc.tables:
        rows = list(table.rows)
        if not rows: continue
        sap_cols = {ci for ci, cell in enumerate(rows[0].cells)
                    if cell.text.strip().lower() in _TABLE_CODE_HEADERS}
        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row.cells):
                t = cell.text.strip()
                if not t: continue
                if ci in sap_cols and ri > 0:
                    cand = _strip_prefix(t.upper().replace('-', '_'))
                    if cand not in _TABLE_IGNORE and _is_valid_tcode(cand):
                        tcodes.add(cand)
                else:
                    tcodes |= _extract_tcodes_from_text(t)
    return sorted(tcodes)


# ---------------------------------------------------------------------------
# Prompt parsing — derives focus mode and instruction hints from PROMPT
# ---------------------------------------------------------------------------

def parse_prompt(prompt: str) -> dict:
    """
    Parse the PROMPT string into structured instructions:
      focus_mode    : what content to compare ('all', 'sap_only', 'steps_only',
                      'activities_only', 'sap_and_steps')
      hint_text     : cleaned text injected into TF-IDF to guide cluster boundaries
      raw           : original prompt for display
    """
    p = prompt.lower().strip()

    # ── Focus mode ────────────────────────────────────────────────────────────
    if re.search(r'compare.{0,20}only.{0,20}sap.{0,20}(code|transaction)', p) or \
       re.search(r'only.{0,10}sap.{0,10}(code|transaction)', p):
        focus = 'sap_only'
    elif re.search(r'compare.{0,20}sap.{0,30}(and|with|together).{0,30}(process|step)', p) or \
         re.search(r'compare.{0,20}(process|step).{0,30}(and|with|together).{0,30}sap', p) or \
         re.search(r'sap.{0,20}(code|transaction).{0,20}and.{0,20}process.{0,20}step', p):
        focus = 'sap_and_steps'
    elif re.search(r'compare.{0,20}only.{0,20}(detailed\s+)?process.{0,10}step', p) or \
         re.search(r'only.{0,10}(detailed\s+)?process.{0,10}step', p):
        focus = 'steps_only'
    elif re.search(r'compare.{0,20}only.{0,20}activit', p) or \
         re.search(r'only.{0,10}activit', p):
        focus = 'activities_only'
    elif re.search(r'(focus|compare).{0,30}(process\s+step|detailed\s+step)', p) or \
         re.search(r'process\s+step.{0,20}(similar|find|focus|group)', p) or \
         'process steps' in p:
        focus = 'steps_only'
    else:
        focus = 'all'

    # ── Extract grouping rules from prompt ───────────────────────────────────
    # Split into sentences and keep instruction-style ones
    sentences = re.split(r'[.\n]', p)
    rules = []
    for s in sentences:
        s = s.strip()
        if len(s) < 8:
            continue
        # Skip pure focus-mode sentences
        if re.match(r'(compare|focus on)\s', s, re.IGNORECASE):
            continue
        rules.append(s)

    hint_text = ' . '.join(rules)

    return {
        'focus_mode': focus,
        'hint_text' : hint_text,
        'rules'     : rules,
        'raw'       : prompt.strip(),
    }


# ---------------------------------------------------------------------------
# Section-aware text extraction
# ---------------------------------------------------------------------------

def is_heading(para) -> bool:
    style_name = para.style.name.lower() if para.style else ''
    return (style_name.startswith('heading') or style_name.startswith('title')
            or (para.runs and all(r.bold for r in para.runs if r.text.strip())))


def heading_matches_target(text: str) -> bool:
    h = text.strip().upper()
    return any(t.upper() in h or h in t.upper() for t in TARGET_SECTIONS)


def extract_section_text(filepath: str) -> tuple:
    doc = Document(filepath)
    para_data = [(is_heading(p), p.text.strip()) for p in doc.paragraphs if p.text.strip()]
    table_texts = [cell.text.strip() for table in doc.tables
                   for row in table.rows for cell in row.cells if cell.text.strip()]
    full_text = ' '.join(t for _, t in para_data) + ' ' + ' '.join(table_texts)

    section_parts, sections_found = [], []
    capturing = False
    for is_hdg, text in para_data:
        if is_hdg:
            capturing = heading_matches_target(text)
            if capturing: sections_found.append(text)
        elif capturing:
            section_parts.append(text)

    # Fallback: ALL-CAPS line matching
    if not section_parts:
        capturing = False
        for is_hdg, text in para_data:
            looks_hdg = is_hdg or (len(text.split()) <= 8 and text == text.upper())
            is_target  = any(t.upper() in text.upper() for t in TARGET_SECTIONS)
            if looks_hdg:
                capturing = is_target
                if capturing and text not in sections_found:
                    sections_found.append(text)
            elif capturing:
                section_parts.append(text)

    return ' '.join(section_parts).strip(), sections_found, full_text


def _extract_named_section(filepath: str, section_name: str) -> str:
    try:
        doc, parts, capturing = Document(filepath), [], False
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text: continue
            is_hdg = is_heading(para) or (len(text.split()) <= 8 and text == text.upper())
            if is_hdg:
                capturing = section_name.upper() in text.upper()
            elif capturing:
                parts.append(text)
        return ' '.join(parts).strip()
    except Exception:
        return ''


def _extract_step_sentences(filepath: str) -> list:
    """
    Extract individual process step sentences from DETAILED PROCESS STEPS section.
    Splits numbered steps (1. 2. 3.) and bullet points into separate sentences.
    These are used to show matching steps in the Reason column.
    """
    raw = _extract_named_section(filepath, 'DETAILED PROCESS STEPS')
    if not raw:
        # Fallback: try any paragraph that looks like a numbered step
        try:
            doc = Document(filepath)
            raw = ' '.join(p.text.strip() for p in doc.paragraphs
                          if re.match(r'^\d+[.)\s]', p.text.strip()))
        except Exception:
            return []

    # Split on numbered steps: "1.", "2.", "Step 1:", bullets
    sentences = re.split(r'(?:^|\s+)(?:\d+[.):]|Step\s+\d+[.):])\s+', raw, flags=re.IGNORECASE)
    sentences = [s.strip().rstrip('.,;') for s in sentences if len(s.strip().split()) >= 3]
    return sentences[:20]  # cap at 20 steps


def apply_focus(focus_mode: str, section_text: str, full_text: str,
                sap_codes: list, filepath: str) -> str:
    if focus_mode == 'sap_only':
        return ' '.join(sap_codes) if sap_codes else section_text or full_text
    if focus_mode == 'sap_and_steps':
        steps = _extract_named_section(filepath, 'DETAILED PROCESS STEPS')
        return (' '.join(sap_codes) + ' ' + (steps or section_text)).strip() or full_text
    if focus_mode == 'steps_only':
        return _extract_named_section(filepath, 'DETAILED PROCESS STEPS') or section_text or full_text
    if focus_mode == 'activities_only':
        return _extract_named_section(filepath, 'SOP ACTIVITIES OVERVIEW') or section_text or full_text
    return section_text or full_text


def clean_text(text: str, strip_regions: bool = True) -> str:
    """Clean text for TF-IDF. Strips region noise by default so regional
    variants of the same process compare equally."""
    if strip_regions:
        text = strip_region_noise(text)
    text = text.lower()
    text = re.sub(r'[^a-z\s]', ' ', text)
    return re.sub(r'\s+', ' ', text).strip()


# ---------------------------------------------------------------------------
# Synonym map — normalises variant words for better semantic matching
# ---------------------------------------------------------------------------
SYNONYM_MAP = {
    'bill':'invoice',       'bills':'invoices',       'billing':'invoicing',
    'supplier':'vendor',    'suppliers':'vendors',    'creditor':'vendor',
    'disposal':'retirement','write-off':'retirement', 'scrapping':'retirement',
    'purchase':'acquisition','entry':'posting',       'entries':'postings',
    'booking':'posting',    'bookings':'postings',    'journal':'ledger',
    'wage':'salary',        'wages':'salaries',       'compensation':'salary',
    'headcount':'employee', 'staff':'employee',       'workforce':'employee',
    'wbs':'project',        'initiative':'project',   'verify':'validate',
    'verification':'validation','check':'review',     'confirm':'approve',
    'confirmation':'approval','authorise':'approve',  'authorization':'approval',
}


def expand_synonyms(text: str) -> str:
    for phrase, rep in SYNONYM_MAP.items():
        if ' ' in phrase:
            text = re.sub(re.escape(phrase), rep, text, flags=re.IGNORECASE)
    words = text.lower().split()
    return ' '.join(SYNONYM_MAP.get(w, w) for w in words)


# ---------------------------------------------------------------------------
# REGION TERMS — words stripped before comparing so regional variants of the
# same process are grouped together.
# Add country names, city names, region codes, or any location-specific terms.
# ---------------------------------------------------------------------------
REGION_TERMS = [
    # Countries
    'india', 'indian', 'uk', 'united kingdom', 'britain', 'british',
    'usa', 'us', 'america', 'american', 'united states',
    'germany', 'german', 'france', 'french', 'china', 'chinese',
    'australia', 'australian', 'singapore', 'japan', 'japanese',
    'canada', 'canadian', 'brazil', 'brazilian', 'mexico', 'mexican',
    'netherlands', 'dutch', 'sweden', 'swedish', 'norway', 'norwegian',
    'uae', 'emirates', 'gulf', 'apac', 'emea', 'latam', 'amer',
    # Regions / zones
    'north', 'south', 'east', 'west', 'central', 'global', 'local',
    'regional', 'international', 'domestic', 'overseas', 'offshore',
    'onshore', 'nearshore',
    # Currency / locale noise
    'gbp', 'usd', 'eur', 'inr', 'sgd', 'aud', 'cad',
    'gst', 'vat', 'tds', 'withholding',
    # Common regional org suffixes
    'ltd', 'llc', 'inc', 'plc', 'gmbh', 'pvt', 'pte',
]

# Compiled once at module load for performance
_REGION_PATTERN = re.compile(
    r'\b(' + '|'.join(re.escape(t) for t in sorted(REGION_TERMS, key=len, reverse=True)) + r')\b',
    re.IGNORECASE
)


def strip_region_noise(text: str) -> str:
    """Remove region/country/locale terms so regional variants compare on process only."""
    text = _REGION_PATTERN.sub(' ', text)
    return re.sub(r'\s+', ' ', text).strip()


# ---------------------------------------------------------------------------
# Model loading (offline enforced)
# ---------------------------------------------------------------------------

_MODEL_CACHE = {}   # cache model so it loads only once per run


def _load_model():
    """Load sentence-transformers in offline mode. Returns model or None (cached)."""
    if 'model' in _MODEL_CACHE:
        return _MODEL_CACHE['model']
    try:
        from sentence_transformers import SentenceTransformer
        os.environ['TRANSFORMERS_OFFLINE'] = '1'
        os.environ['HF_DATASETS_OFFLINE']  = '1'
        model = SentenceTransformer(EMBEDDING_MODEL, local_files_only=True)
        _MODEL_CACHE['model'] = model
        return model
    except ImportError:
        _MODEL_CACHE['model'] = None
        return None
    except Exception as e:
        if any(k in str(e).lower() for k in ['local_files_only', 'no such file', 'not found', 'snapshot']):
            print(f"\n  [INFO] Model '{EMBEDDING_MODEL}' not found in local cache.")
            print(f"  [INFO] Run once (internet needed only for this):")
            print(f"  [INFO]   pip install sentence-transformers")
            print(f"  [INFO]   python -c \"from sentence_transformers import SentenceTransformer; SentenceTransformer('{EMBEDDING_MODEL}')\"")
            print(f"  [INFO] After that, runs fully offline forever. Using TF-IDF fallback now.\n")
        else:
            print(f"  [WARN] Model load failed: {e}. Using TF-IDF fallback.")
        return None


# ---------------------------------------------------------------------------
# Clustering — semantic similarity analysis
# ---------------------------------------------------------------------------

def cluster_documents(texts: list, n_clusters: int, hint_text: str):
    """
    Semantic Similarity Analysis:
      Tier 1 (sentence-transformers): neural embeddings → cosine similarity
      Tier 2 (TF-IDF fallback):       synonym-expanded vectors → cosine similarity
      Both use Agglomerative Clustering (merges most-similar pairs, better than KMeans)
    """
    n_clusters = max(2, min(n_clusters, len(texts) - 1))

    model = _load_model()

    # ── Inject prompt hint as similarity bias ───────────────────────────────
    # The hint is appended to each doc's text. This boosts vocabulary from
    # the prompt so docs sharing those terms score closer to each other.
    # We repeat the hint proportionally to give it meaningful weight without
    # overwhelming the document's own content.
    if hint_text.strip():
        hint_clean  = expand_synonyms(clean_text(hint_text, strip_regions=False))
        hint_repeat = (hint_clean + ' ') * 2   # repeat 2x for moderate boost
        augmented   = [t + ' ' + hint_repeat for t in texts]
    else:
        augmented = texts[:]

    if model is not None:
        print(f"  [Semantic engine: sentence-transformers / {EMBEDDING_MODEL} — fully offline]")
        embeddings = model.encode(augmented, convert_to_numpy=True,
                                  show_progress_bar=False, normalize_embeddings=True)
        sim_matrix = np.clip(np.dot(embeddings, embeddings.T), -1, 1)
    else:
        print("  [Semantic engine: TF-IDF cosine similarity (run one-time setup for better accuracy)]")
        expanded = [expand_synonyms(t) for t in augmented]
        vec = TfidfVectorizer(stop_words='english', max_features=5000,
                              ngram_range=(1, 2), min_df=1, sublinear_tf=True)
        mat = vec.fit_transform(expanded)
        sim_matrix = sklearn_cosine(mat)

    dist_matrix = np.clip(1 - sim_matrix, 0, None)
    labels = AgglomerativeClustering(
        n_clusters=n_clusters, metric='precomputed', linkage='average'
    ).fit_predict(dist_matrix)

    # Always build TF-IDF for keyword extraction
    expanded_kw = [expand_synonyms(t) for t in texts]
    vec_kw = TfidfVectorizer(stop_words='english', max_features=5000,
                             ngram_range=(1, 2), min_df=1, sublinear_tf=True)
    tfidf_kw = vec_kw.fit_transform(expanded_kw)

    return labels, vec_kw, tfidf_kw


# ---------------------------------------------------------------------------
# Group naming and reasoning
# ---------------------------------------------------------------------------

NAME_EXCLUSIONS = {
    'manager','officer','director','head','lead','owner','coordinator',
    'analyst','specialist','controller','executive','assistant','associate',
    'team','staff','personnel','user','users','member','members',
    'january','february','march','april','june','july','august',
    'september','october','november','december','monthly','weekly','daily',
    'report','reports','form','forms','email','document','documents',
    'system','systems','data','information','details','record','records',
}

PROCESS_VERBS = {
    'posting','payment','approval','settlement','reconciliation','processing',
    'capitalisation','capitalization','depreciation','retirement','allocation',
    'budgeting','onboarding','invoicing','procurement','reporting','closing',
    'review','validation','verification','calculation','execution','creation',
    'submission','confirmation','adjustment','transfer','accrual','clearance',
}

PROCESS_NOUNS = {
    'invoice','invoices','vendor','vendors','asset','assets','payroll',
    'salary','salaries','journal','journals','ledger','account','accounts',
    'budget','cost','costs','project','projects','employee','employees',
    'payment','payments','order','orders','purchase','expense','expenses',
    'depreciation','capital','period','balance','balances','bank','tax',
    'revenue','profit','loss','fixed','current','accrual','liability',
}


def looks_like_tcode(word: str) -> bool:
    w = word.upper()
    return bool(re.match(r'^\d+$', w) or re.match(r'^S_ALR', w) or
                re.match(r'^[ZY][A-Z0-9_]{2,}$', w) or
                re.match(r'^[A-Z]{1,3}\d+[A-Z]?$', w) or ('_' in w and len(w) > 6))


def is_excluded_word(word: str) -> bool:
    w = word.lower()
    return (w in STOPWORDS or w in NAME_EXCLUSIONS or looks_like_tcode(word)
            or not word.isalpha() or len(word) <= 3)


def top_cluster_keywords(vectorizer, cluster_id: int, labels, tfidf_matrix, n=30) -> list:
    feature_names = vectorizer.get_feature_names_out()
    members = [i for i, l in enumerate(labels) if l == cluster_id]
    if not members: return []
    centroid = np.asarray(tfidf_matrix[members].mean(axis=0)).flatten()
    keywords = []
    for idx in centroid.argsort()[::-1]:
        w = feature_names[idx]
        if not is_excluded_word(w):
            keywords.append(w)
        if len(keywords) >= n:
            break
    return keywords


def make_group_name(keywords: list, cluster_id: int) -> str:
    if not keywords:
        return f"Process Group {cluster_id + 1}"
    verbs  = [w for w in keywords if w.lower() in PROCESS_VERBS]
    nouns  = [w for w in keywords if w.lower() in PROCESS_NOUNS]
    parts  = []
    if verbs and nouns:
        parts.append(f"{nouns[0].title()} {verbs[0].title()}")
        n2 = nouns[1].title() if len(nouns) > 1 else nouns[0].title()
        v2 = verbs[1].title() if len(verbs) > 1 else verbs[0].title()
        if f"{n2} {v2}" != parts[0]: parts.append(f"{n2} {v2}")
    elif verbs: parts = [w.title() for w in verbs[:3]]
    elif nouns: parts = [w.title() for w in nouns[:3]]
    else:       parts = [w.title() for w in keywords[:3]]
    return " & ".join(parts[:2]) if len(parts) >= 2 else parts[0]


def doc_top_keywords(doc_vector, vectorizer, n=DOC_REASON_KEYWORDS) -> list:
    feature_names = vectorizer.get_feature_names_out()
    arr = np.asarray(doc_vector.todense()).flatten()
    keywords = []
    for idx in arr.argsort()[::-1]:
        if arr[idx] == 0: break
        w = feature_names[idx]
        if not is_excluded_word(w):
            keywords.append(w)
        if len(keywords) >= n: break
    return keywords


def extract_step_sentences(filepath: str, section_text: str = '', full_text: str = '') -> list:
    """
    Extract individual process steps from the DETAILED PROCESS STEPS section.
    Each paragraph or numbered item becomes one step.
    Region noise is stripped so regional variants compare purely on process action.
    """
    steps = []

    # First try: read each paragraph individually from the docx (preserves step boundaries)
    try:
        doc = Document(filepath)
        capturing = False
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            is_hdg = is_heading(para) or (len(text.split()) <= 8 and text == text.upper())
            if is_hdg:
                capturing = 'DETAILED PROCESS STEPS' in text.upper() or                             any(t.upper() in text.upper() for t in TARGET_SECTIONS
                                if 'STEP' in t.upper())
            elif capturing:
                # Strip step number prefix: "1.", "1)", "Step 1:", "Step 1 -"
                clean = re.sub(r'^\s*(\d+[\.):]?\s*|Step\s*\d+\s*[:\-]?\s*)', '',
                               text, flags=re.IGNORECASE).strip()
                clean = strip_region_noise(clean)
                clean = re.sub(r'\s+', ' ', clean).strip()
                # Keep meaningful steps only (not sub-headers, not too short)
                if 4 <= len(clean.split()) <= 50 and not clean.isupper():
                    steps.append(clean)
    except Exception:
        pass

    # Fallback: split section_text on numbered patterns if paragraph read failed
    if not steps:
        source = section_text or full_text
        # Split on numbered step starts: "1. " "2. " "Step 3: "
        parts = re.split(r'(?<=\n)(?:\d+[\.):]\s+|Step\s+\d+[:\-]\s+)', source, flags=re.IGNORECASE)
        for part in parts:
            clean = strip_region_noise(part.strip())
            # Take first sentence of each part
            first_sentence = re.split(r'(?<=[.!?])\s+', clean)[0].strip()
            first_sentence = re.sub(r'\s+', ' ', first_sentence)
            if 4 <= len(first_sentence.split()) <= 50:
                steps.append(first_sentence)

    return steps[:25]   # cap at 25 steps per doc


def _embed(texts: list, model) -> np.ndarray:
    """Batch encode texts into normalised embedding vectors."""
    cleaned = [expand_synonyms(strip_region_noise(s)) for s in texts if s.strip()]
    if not cleaned:
        return np.zeros((0, 1))
    return model.encode(cleaned, convert_to_numpy=True,
                        show_progress_bar=False, normalize_embeddings=True)


def _jaccard(s1: str, s2: str) -> float:
    """Word-overlap Jaccard after synonym expansion."""
    w1 = set(expand_synonyms(clean_text(s1)).split()) - set(STOPWORDS)
    w2 = set(expand_synonyms(clean_text(s2)).split()) - set(STOPWORDS)
    return len(w1 & w2) / len(w1 | w2) if w1 and w2 and (w1 | w2) else 0.0


def step_similarity_score(steps_a: list, steps_b: list) -> float:
    """
    Overall similarity between two documents' step lists.

    WITH model   : embedding cosine similarity per step pair — understands
                   meaning, so "park invoice FBV3" ≈ "hold bill for review"
    WITHOUT model: Jaccard word overlap after synonym expansion
    """
    if not steps_a or not steps_b:
        return 0.0

    model = _load_model()
    if model is not None:
        emb_a = _embed(steps_a, model)
        emb_b = _embed(steps_b, model)
        if emb_a.size == 0 or emb_b.size == 0:
            return 0.0
        sim = np.dot(emb_a, emb_b.T)        # (len_a, len_b) cosine matrix
        return float(sim.max(axis=1).mean()) # average best-match per step in A
    else:
        scores = [max((_jaccard(sa, sb) for sb in steps_b), default=0.0)
                  for sa in steps_a]
        return sum(scores) / len(scores) if scores else 0.0


def find_matching_steps(doc_steps: list, other_docs_steps: list,
                        top_n: int = 5, threshold: float = 0.35) -> list:
    """
    Find which steps in this document are most similar to steps from
    all other documents in the same cluster.

    WITH model   : embedding cosine — "Verify MwSt" matches "Validate VAT"
    WITHOUT model: Jaccard word overlap with lower threshold

    Returns list of (score, this_doc_step, best_matching_step_from_cluster).
    """
    if not doc_steps or not other_docs_steps:
        return []

    all_other = [s for steps in other_docs_steps for s in steps]
    if not all_other:
        return []

    model = _load_model()
    matches = []

    if model is not None:
        emb_doc   = _embed(doc_steps, model)
        emb_other = _embed(all_other,  model)
        if emb_doc.size == 0 or emb_other.size == 0:
            return []
        sim = np.dot(emb_doc, emb_other.T)    # (n_doc_steps, n_other_steps)
        for i, step in enumerate(doc_steps):
            j    = int(sim[i].argmax())
            score = float(sim[i, j])
            if score >= threshold:
                matches.append((score, step, all_other[j]))
    else:
        jac_threshold = threshold * 0.4       # Jaccard needs lower bar
        for step in doc_steps:
            best_score, best_match = 0.0, ""
            for other in all_other:
                s = _jaccard(step, other)
                if s > best_score:
                    best_score, best_match = s, other
            if best_score >= jac_threshold:
                matches.append((best_score, step, best_match))

    # Sort by score, deduplicate on doc_step
    matches.sort(key=lambda x: -x[0])
    seen, out = set(), []
    for score, ds, ms in matches:
        if ds not in seen:
            seen.add(ds)
            out.append((score, ds, ms))
    return out[:top_n]


def make_reason(doc_kws, cluster_kws, group_name, sections_found,
                used_fallback, focus_mode, doc_steps=None,
                cluster_steps=None,   # list of step lists from OTHER docs in cluster
                prompt_rules=None) -> str:
    lines = []

    # ── 1. What was compared ─────────────────────────────────────────────────
    if focus_mode == 'sap_only':
        lines.append("Basis: SAP transaction code similarity")
    elif focus_mode == 'sap_and_steps':
        lines.append("Basis: SAP codes + process steps (region-neutral)")
    elif focus_mode == 'steps_only':
        lines.append("Basis: Detailed Process Steps (region-neutral)")
    elif sections_found:
        lines.append("Basis: " + ' & '.join(s.title() for s in sections_found[:2]))
    elif used_fallback:
        lines.append("Basis: Full document content (target sections not found)")

    # ── 2. Prompt rules that apply ───────────────────────────────────────────
    if prompt_rules:
        # Show only rules that seem relevant to this group
        grp_lower = group_name.lower()
        relevant = [r for r in prompt_rules
                    if any(w in grp_lower for w in r.lower().split()
                           if len(w) > 4 and w not in STOPWORDS)]
        if not relevant:
            relevant = prompt_rules[:1]   # show at least first rule
        for rule in relevant[:2]:
            lines.append(f"Prompt rule applied: {rule.strip().rstrip('.')}")

    # ── 3. Matching process steps — pairwise evidence ───────────────────────
    # cluster_steps is a list of step lists from all other docs in same cluster
    if doc_steps and cluster_steps:
        other_steps_lists = cluster_steps if isinstance(cluster_steps[0], list) else [cluster_steps]
        matches = find_matching_steps(doc_steps, other_steps_lists, top_n=5)
        model   = _load_model()
        engine  = "semantic embedding" if model is not None else "keyword overlap"
        if matches:
            lines.append(f"Similar process steps ({engine}):")
            for score, this_step, matched_step in matches:
                first = re.split(r'[.!?]\s+', this_step.strip())[0].strip()
                first = re.sub(r'\s+', ' ', first)
                disp = first[:100] + ("..." if len(first) > 100 else "")
                pct  = int(score * 100)
                lines.append(f"  • {disp}  [{pct}%]")
        else:
            lines.append("⚠ No strongly matching steps found — may be incorrectly grouped. Review manually.")
            lines.append("Top steps from this document:")
            for step in doc_steps[:3]:
                first = re.split(r'[.!?]\s+', step.strip())[0].strip()[:90]
                lines.append(f"  • {first}")
    elif doc_steps:
        lines.append("Process steps extracted (only document in group):")
        for step in doc_steps[:4]:
            lines.append(f"  • {step[:85]}")

    # ── 4. Keyword fallback if no steps extracted at all ─────────────────────
    if not doc_steps:
        shared_v = [k for k in doc_kws if k.lower() in PROCESS_VERBS and k in cluster_kws]
        shared_n = [k for k in doc_kws if k.lower() in PROCESS_NOUNS and k in cluster_kws]
        if shared_v: lines.append("Shared activities: " + ', '.join(w.title() for w in shared_v[:3]))
        if shared_n: lines.append("Shared subjects: " + ', '.join(w.title() for w in shared_n[:3]))
        if not shared_v and not shared_n:
            lines.append("Note: No process steps found — check if TARGET_SECTIONS headings match document.")

    return '\n'.join(lines) if lines else f"Similar process flow to {group_name}."


def optimal_clusters(n_docs, requested=None):
    if requested: return min(requested, n_docs)
    return max(2, min(8, n_docs - 1, round(n_docs / 3)))


# ---------------------------------------------------------------------------
# Excel report
# ---------------------------------------------------------------------------

def build_report(groups: dict, output_path: str, prompt_config: dict):
    wb = Workbook()
    ws = wb.active
    ws.title = "SOP Groups"

    hdr_font  = Font(name='Arial', bold=True, color='FFFFFF', size=11)
    hdr_fill  = PatternFill('solid', start_color='1F4E79')
    hdr_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    grp_font  = Font(name='Arial', bold=True, color='FFFFFF', size=10)
    doc_font  = Font(name='Arial', size=10)
    rsn_font  = Font(name='Arial', size=10, italic=True)
    tcd_font  = Font(name='Courier New', size=10, color='1F4E79')
    top_align = Alignment(horizontal='left', vertical='top', wrap_text=True)
    mid_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin_s = Side(style='thin', color='BDD7EE')
    med_s  = Side(style='medium', color='1F4E79')
    thin_b = Border(left=thin_s, right=thin_s, top=thin_s, bottom=thin_s)
    med_b  = Border(left=med_s,  right=med_s,  top=med_s,  bottom=med_s)

    PALETTE   = ['1F4E79','2E75B6','70AD47','ED7D31','FFC000','5B9BD5','A9D18E','F4B183','FFD966','9DC3E6']
    ROW_FILLS = [PatternFill('solid', start_color=c) for c in
                 ['EBF3FB','EFF7E6','FFF2CC','FCE4D6','DAEEF3']]

    # Banner — show active prompt
    ws.merge_cells('A1:D1')
    c = ws['A1']
    c.value = f"Prompt: {prompt_config['raw']}  |  Focus: {prompt_config['focus_mode']}  |  Sections: {', '.join(TARGET_SECTIONS)}"
    c.font      = Font(name='Arial', bold=True, color='1F4E79', size=9)
    c.fill      = PatternFill('solid', start_color='DEEAF1')
    c.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
    c.border    = med_b
    ws.row_dimensions[1].height = 30

    for col, h in enumerate(['Group Name','Document Name','Reason for Grouping','SAP Transaction Codes'], 1):
        c = ws.cell(row=2, column=col, value=h)
        c.font = hdr_font; c.fill = hdr_fill
        c.alignment = hdr_align; c.border = med_b
    ws.row_dimensions[2].height = 28

    REF_GRP_FILL = PatternFill('solid', start_color='808080')  # grey for reference docs
    REF_ROW_FILL = PatternFill('solid', start_color='F2F2F2')

    row = 3
    # Sort: process groups first alphabetically, Reference Documents last
    sorted_groups = sorted(
        groups.items(),
        key=lambda x: ('~' + x[0]) if x[0] == 'Reference Documents' else x[0]
    )
    for g_idx, (group_name, docs) in enumerate(sorted_groups):
        is_ref    = (group_name == 'Reference Documents')
        grp_fill  = REF_GRP_FILL if is_ref else PatternFill('solid', start_color=PALETTE[g_idx % len(PALETTE)])
        row_fill  = REF_ROW_FILL if is_ref else ROW_FILLS[g_idx % len(ROW_FILLS)]
        start_row, n = row, len(docs)

        for rel_path, reason, sap_codes in docs:
            ws.cell(row=row, column=2, value=rel_path).font = doc_font
            ws.cell(row=row, column=2).fill = row_fill
            ws.cell(row=row, column=2).alignment = top_align
            ws.cell(row=row, column=2).border = thin_b

            ws.cell(row=row, column=3, value=reason).font = rsn_font
            ws.cell(row=row, column=3).fill = row_fill
            ws.cell(row=row, column=3).alignment = top_align
            ws.cell(row=row, column=3).border = thin_b

            codes_str = ', '.join(sap_codes) if sap_codes else '—'
            ws.cell(row=row, column=4, value=codes_str).font = tcd_font
            ws.cell(row=row, column=4).fill = row_fill
            ws.cell(row=row, column=4).alignment = top_align
            ws.cell(row=row, column=4).border = thin_b

            ws.row_dimensions[row].height = max(45, 15 * len(str(reason).split('\n')))
            row += 1

        c = ws.cell(row=start_row, column=1, value=group_name)
        c.font = grp_font; c.fill = grp_fill
        c.alignment = mid_align; c.border = med_b
        if n > 1:
            ws.merge_cells(start_row=start_row, start_column=1, end_row=row-1, end_column=1)
            c = ws.cell(row=start_row, column=1)
            c.font = grp_font; c.fill = grp_fill
            c.alignment = mid_align; c.border = med_b

    ws.column_dimensions['A'].width = 28
    ws.column_dimensions['B'].width = 45
    ws.column_dimensions['C'].width = 68
    ws.column_dimensions['D'].width = 38
    ws.freeze_panes = 'B3'
    wb.save(output_path)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--groups', '-g', type=int, default=None)
    args = parser.parse_args()

    script_dir  = os.path.dirname(os.path.abspath(__file__))
    input_dir   = os.path.join(script_dir, 'input')
    output_dir  = os.path.join(script_dir, 'output')
    output_file = os.path.join(output_dir, 'SOP_Groups.xlsx')
    os.makedirs(input_dir, exist_ok=True)
    os.makedirs(output_dir, exist_ok=True)

    prompt_config = parse_prompt(PROMPT)
    focus_mode    = prompt_config['focus_mode']
    hint_text     = prompt_config['hint_text']

    print(f"Input  folder : {input_dir}")
    print(f"Output folder : {output_dir}")
    print(f"Target sections: {TARGET_SECTIONS}")
    print(f"Focus mode     : {focus_mode}")
    print(f"Prompt         : {prompt_config['raw'][:120]}...")
    print()

    all_docx = sorted([
        os.path.join(dp, f)
        for dp, _, fs in os.walk(input_dir)
        for f in fs if f.lower().endswith('.docx') and not f.startswith('~$')
    ])

    if not all_docx:
        raise SystemExit(f"No .docx files found in '{input_dir}'.")

    print(f"Found {len(all_docx)} document(s). Extracting content...\n")

    rel_paths, cluster_texts, meta = [], [], []

    for filepath in all_docx:
        rel = os.path.relpath(filepath, input_dir)
        try:
            section_text, sections_found, full_text = extract_section_text(filepath)
            sap_codes = extract_sap_codes(filepath)
        except Exception as e:
            print(f"  [WARN] Could not read '{rel}': {e}")
            continue

        focused = apply_focus(focus_mode, section_text, full_text, sap_codes, filepath)
        used_fallback = (len(section_text.split()) < 10)

        cleaned = clean_text(focused)
        word_count = len(cleaned.split())

        # Documents with fewer than MIN_PROCESS_WORDS in their process content
        # are classified as "Reference Documents" — short guides, quick refs,
        # checklists — rather than forced into a process group.
        # Priority: process steps similarity first, region is secondary noise.
        if word_count < MIN_PROCESS_WORDS:
            print(f"  * {rel}  [Reference Document — only {word_count} words in process content]")
            rel_paths.append(rel)
            cluster_texts.append(cleaned if cleaned else 'reference document quick guide')
            meta.append({
                'rel': rel, 'sections_found': sections_found,
                'used_fallback': used_fallback, 'sap_codes': sap_codes,
                'is_reference': True, 'step_sentences': [],
            })
            continue

        # Extract individual step sentences for reason generation
        step_sentences = extract_step_sentences(filepath, section_text, full_text)

        if focus_mode == 'sap_only':
            print(f"  + {rel}  [SAP codes: {', '.join(sap_codes) or 'none'}]")
        elif not used_fallback:
            print(f"  + {rel}  [sections: {', '.join(sections_found)}]")
        else:
            print(f"  ~ {rel}  [FALLBACK: sections not found, using full text]")

        rel_paths.append(rel)
        cluster_texts.append(cleaned)
        meta.append({
            'rel': rel, 'sections_found': sections_found,
            'used_fallback': used_fallback, 'sap_codes': sap_codes,
            'is_reference': False, 'step_sentences': step_sentences,
        })

    n_docs = len(rel_paths)
    if n_docs < MIN_DOCS:
        raise SystemExit(f"\nNeed at least {MIN_DOCS} documents. Found {n_docs}.")

    n_clusters = optimal_clusters(n_docs, args.groups)
    print(f"\nClustering {n_docs} document(s) into {n_clusters} group(s)...")

    labels, vectorizer, tfidf_matrix = cluster_documents(cluster_texts, n_clusters, hint_text)

    cluster_kws   = {i: top_cluster_keywords(vectorizer, i, labels, tfidf_matrix)
                     for i in range(n_clusters)}
    cluster_names = {i: make_group_name(cluster_kws[i], i) for i in range(n_clusters)}

    # Pre-build cluster step pools: aggregate all step sentences per cluster
    cluster_steps_pool = defaultdict(list)
    for idx, label in enumerate(labels):
        m = meta[idx]
        if not m.get('is_reference'):
            cluster_steps_pool[label].extend(m.get('step_sentences', []))

    groups = {}
    for idx, label in enumerate(labels):
        m          = meta[idx]

        # Reference documents go to their own fixed group regardless of cluster
        if m.get('is_reference'):
            group_name = 'Reference Documents'
            reason     = 'Short document (quick reference, checklist, or summary) — insufficient process steps for similarity matching.'
            groups.setdefault(group_name, []).append((m['rel'], reason, m['sap_codes']))
            print(f"  [Reference Documents]  {m['rel']}  (reference doc)")
            continue

        group_name   = cluster_names[label]
        doc_kws      = doc_top_keywords(tfidf_matrix[idx], vectorizer)
        doc_steps    = m.get('step_sentences', [])
        # Pass steps from OTHER docs in cluster (not this doc — pairwise comparison)
        other_steps  = [meta[j]['step_sentences'] for j, lbl in enumerate(labels)
                        if lbl == label and j != idx and not meta[j].get('is_reference')]
        reason       = make_reason(doc_kws, cluster_kws[label], group_name,
                                   m['sections_found'], m['used_fallback'], focus_mode,
                                   doc_steps=doc_steps, cluster_steps=other_steps,
                                   prompt_rules=prompt_config.get('rules', []))
        groups.setdefault(group_name, []).append((m['rel'], reason, m['sap_codes']))
        print(f"  [{group_name}]  {m['rel']}")

    build_report(groups, output_file, prompt_config)
    print(f"\nReport saved: {output_file}")
    print(f"  {n_clusters} group(s), {n_docs} document(s)\n")
    for name, docs in sorted(groups.items()):
        print(f"  {name}: {len(docs)} doc(s)")


if __name__ == '__main__':
    main()
