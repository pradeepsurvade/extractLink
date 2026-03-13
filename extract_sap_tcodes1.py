"""
SAP Transaction Code Extractor
================================
Scans a folder of Word (.docx) documents, extracts SAP transaction codes,
and outputs an Excel file with:
  Column A: Transaction Code
  Column B: Document name(s) containing that code
  Column C: Source sentence from which the code was extracted

SAP T-Code patterns detected (based on screenshot logic):
  - S_ALR_XXXXXXX      (e.g. S_ALR_87013611, S_ALR_87011964)
  - ZXXX_XXX_XXX       (custom Z-codes, e.g. ZRTR_GL_PARK_GL)
  - YXXX_XXX_XXX       (custom Y-codes, e.g. YRTR_ASSET_BALANCES)
  - Standard codes     (e.g. CJ20N, CJ88, FB01, MIGO, ME21N)

Usage:
  python extract_sap_tcodes.py --folder /path/to/docs --output result.xlsx
"""

import os
import re
from collections import defaultdict

try:
    from docx import Document
except ImportError:
    raise SystemExit("python-docx not installed. Run: pip install python-docx --break-system-packages")

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
except ImportError:
    raise SystemExit("openpyxl not installed. Run: pip install openpyxl --break-system-packages")


# ── SAP T-Code regex patterns ─────────────────────────────────────────────────
TCODE_PATTERNS = [
    re.compile(r'\bS_ALR_\d{7,8}\b'),
    re.compile(r'\b[ZY][A-Z0-9]{1,15}(?:_[A-Z0-9]{1,20})+\b'),
    re.compile(r'\b[A-Z]{1,4}[0-9]{1,4}[A-Z]?\b'),
]

TRIGGER_PHRASES = [
    r'run\s+',
    r't[-\s]?code[s]?\s*[:\(]?\s*',
    r'tcode[s]?\s*[:\(]?\s*',
    r'transaction\s*[:\(]?\s*',
    r'sap\s+t\s*code\s*',
    r'run\s+the\s+',
]
TRIGGER_RE = re.compile(
    '(' + '|'.join(TRIGGER_PHRASES) + r')([A-Z0-9_]{2,30})',
    re.IGNORECASE
)

EXCLUSIONS = {
    'THE', 'AND', 'FOR', 'RUN', 'SAP', 'CODE', 'WITH', 'FROM', 'INTO',
    'THIS', 'THAT', 'HAVE', 'BEEN', 'WILL', 'ALSO', 'CAN', 'NOT', 'ARE',
    'ALL', 'BUT', 'USE', 'NEW', 'OLD', 'END', 'ADD', 'SET', 'GET', 'PUT',
    'OUT', 'OFF', 'TOP', 'BOX', 'YES', 'NO', 'OK', 'GO', 'DO', 'IF',
}


def is_valid_tcode(candidate: str) -> bool:
    c = candidate.upper().strip()
    if c in EXCLUSIONS or len(c) < 2:
        return False
    if re.match(r'^S_ALR_\d{5,}$', c):
        return True
    if re.match(r'^[ZY][A-Z0-9_]{3,}$', c) and '_' in c:
        return True
    if re.match(r'^[A-Z]{1,4}\d{1,4}[A-Z]?$', c):
        return True
    return False


def extract_tcodes_from_text(text: str) -> list:
    """
    Returns list of (tcode, source_sentence) tuples found in text.
    """
    results = []
    text_stripped = text.strip()
    if not text_stripped:
        return results

    # Strategy 1: trigger phrase → code
    for m in TRIGGER_RE.finditer(text):
        candidate = m.group(2).strip().rstrip('.,;)(')
        if is_valid_tcode(candidate):
            results.append((candidate.upper(), text_stripped))

    # Strategy 2: structural pattern match
    for pattern in TCODE_PATTERNS:
        for m in pattern.finditer(text):
            candidate = m.group(0)
            if is_valid_tcode(candidate):
                # Avoid duplicate tcodes from the same sentence
                already = any(r[0] == candidate.upper() and r[1] == text_stripped for r in results)
                if not already:
                    results.append((candidate.upper(), text_stripped))

    return results


def extract_from_docx(filepath: str) -> list:
    """
    Returns list of (tcode, doc_name, source_sentence) for all matches in a docx.
    """
    doc_name = os.path.basename(filepath)
    doc = Document(filepath)
    matches = []

    for para in doc.paragraphs:
        for tcode, sentence in extract_tcodes_from_text(para.text):
            matches.append((tcode, doc_name, sentence))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for tcode, sentence in extract_tcodes_from_text(cell.text):
                    matches.append((tcode, doc_name, sentence))

    return matches


def scan_folder(folder: str):
    """
    Scan all .docx files in folder.
    Returns:
      - tcode_data: {tcode: [(doc_name, sentence), ...]}
      - others: [doc_names with no t-codes found]
    """
    tcode_data = defaultdict(list)
    others = []

    docx_files = [
        f for f in os.listdir(folder)
        if f.lower().endswith('.docx') and not f.startswith('~$')
    ]

    if not docx_files:
        raise SystemExit(f"No .docx files found in: {folder}")

    print(f"Found {len(docx_files)} Word document(s) to scan...\n")

    for filename in sorted(docx_files):
        path = os.path.join(folder, filename)
        try:
            matches = extract_from_docx(path)
        except Exception as e:
            print(f"  [WARN] Could not read '{filename}': {e}")
            others.append(filename)
            continue

        if matches:
            for tcode, doc_name, sentence in matches:
                tcode_data[tcode].append((doc_name, sentence))
            print(f"  ✔ {filename} → {len(set(m[0] for m in matches))} t-code(s) found")
        else:
            others.append(filename)
            print(f"  ✘ {filename} → no t-codes found (→ Others)")

    return tcode_data, others


def apply_cell_style(cell, font, fill, alignment, border):
    cell.font = font
    cell.fill = fill
    cell.alignment = alignment
    cell.border = border


def build_excel(tcode_data: dict, others: list, output_path: str):
    wb = Workbook()
    ws = wb.active
    ws.title = "SAP T-Codes"

    # ── Styles ────────────────────────────────────────────────────────────────
    hdr_font      = Font(name='Arial', bold=True, color='FFFFFF', size=11)
    hdr_fill      = PatternFill('solid', start_color='1F4E79')
    hdr_align     = Alignment(horizontal='center', vertical='center', wrap_text=True)

    tcode_font    = Font(name='Courier New', bold=True, color='1F4E79', size=10)
    others_font   = Font(name='Courier New', bold=True, color='C00000', size=10)
    doc_font      = Font(name='Arial', size=10)
    sentence_font = Font(name='Arial', size=10, italic=True)
    grey_font     = Font(name='Arial', size=10, italic=True, color='808080')

    mid_align = Alignment(horizontal='left', vertical='center', wrap_text=True)
    top_align = Alignment(horizontal='left', vertical='top',    wrap_text=True)

    thin   = Side(style='thin',   color='BDD7EE')
    medium = Side(style='medium', color='1F4E79')

    def cell_border(top=thin, bottom=thin, left=thin, right=thin):
        return Border(top=top, bottom=bottom, left=left, right=right)

    alt_fill  = PatternFill('solid', start_color='EBF3FB')
    base_fill = PatternFill('solid', start_color='FFFFFF')

    # ── Headers ───────────────────────────────────────────────────────────────
    for col, h in enumerate(['Transaction Code', 'Document Name', 'Source Sentence'], 1):
        c = ws.cell(row=1, column=col, value=h)
        c.font      = hdr_font
        c.fill      = hdr_fill
        c.alignment = hdr_align
        c.border    = cell_border(top=medium, bottom=medium,
                                  left=medium, right=medium)
    ws.row_dimensions[1].height = 32

    # ── Data rows — Column A merged per tcode, B & C one row each entry ───────
    row = 2
    sorted_tcodes = sorted(tcode_data.keys())
    block = 0  # for alternating fill

    for tcode in sorted_tcodes:
        entries   = tcode_data[tcode]   # list of (doc_name, sentence)
        n         = len(entries)
        start_row = row
        fill      = alt_fill if (block % 2 == 0) else base_fill
        block    += 1

        # ── Column B & C: one row per (doc_name, sentence) ──────────────────
        for doc_name, sentence in entries:
            # Col B
            c = ws.cell(row=row, column=2, value=doc_name)
            c.font      = doc_font
            c.fill      = fill
            c.alignment = top_align
            c.border    = cell_border()

            # Col C
            c = ws.cell(row=row, column=3, value=sentence)
            c.font      = sentence_font
            c.fill      = fill
            c.alignment = top_align
            c.border    = cell_border()

            ws.row_dimensions[row].height = 38
            row += 1

        # ── Column A: write tcode once, then merge all rows for this block ──
        end_row = row - 1
        c = ws.cell(row=start_row, column=1, value=tcode)
        c.font      = tcode_font
        c.fill      = fill
        c.alignment = Alignment(horizontal='center', vertical='center',
                                wrap_text=True)
        c.border    = cell_border(top=medium, bottom=medium,
                                  left=medium, right=medium)

        if n > 1:
            ws.merge_cells(start_row=start_row, start_column=1,
                           end_row=end_row,   end_column=1)
            # Re-apply style to merged cell (openpyxl requires this)
            ws.cell(row=start_row, column=1).font      = tcode_font
            ws.cell(row=start_row, column=1).fill      = fill
            ws.cell(row=start_row, column=1).alignment = Alignment(
                horizontal='center', vertical='center', wrap_text=True)
            ws.cell(row=start_row, column=1).border = cell_border(
                top=medium, bottom=medium, left=medium, right=medium)

    # ── Others section ────────────────────────────────────────────────────────
    if others:
        others_fill     = PatternFill('solid', start_color='C00000')
        others_entries  = [(doc, 'No SAP transaction code detected')
                           for doc in sorted(others)]
        n               = len(others_entries)
        start_row       = row

        for doc_name, note in others_entries:
            fill = alt_fill if (row % 2 == 0) else base_fill

            c = ws.cell(row=row, column=2, value=doc_name)
            c.font = doc_font; c.fill = fill
            c.alignment = top_align; c.border = cell_border()

            c = ws.cell(row=row, column=3, value=note)
            c.font = grey_font; c.fill = fill
            c.alignment = top_align; c.border = cell_border()

            ws.row_dimensions[row].height = 30
            row += 1

        end_row = row - 1

        # Col A: "Others" merged
        c = ws.cell(row=start_row, column=1, value='Others')
        c.font      = others_font
        c.fill      = PatternFill('solid', start_color='FFE7E7')
        c.alignment = Alignment(horizontal='center', vertical='center',
                                wrap_text=True)
        c.border    = cell_border(top=medium, bottom=medium,
                                  left=medium, right=medium)

        if n > 1:
            ws.merge_cells(start_row=start_row, start_column=1,
                           end_row=end_row,   end_column=1)
            ws.cell(row=start_row, column=1).font      = others_font
            ws.cell(row=start_row, column=1).fill      = PatternFill('solid', start_color='FFE7E7')
            ws.cell(row=start_row, column=1).alignment = Alignment(
                horizontal='center', vertical='center', wrap_text=True)
            ws.cell(row=start_row, column=1).border = cell_border(
                top=medium, bottom=medium, left=medium, right=medium)

    # ── Column widths & finishing ─────────────────────────────────────────────
    ws.column_dimensions['A'].width = 26
    ws.column_dimensions['B'].width = 38
    ws.column_dimensions['C'].width = 82

    ws.freeze_panes = 'B2'
    ws.auto_filter.ref = f"A1:C1"   # header-only filter (merged col A can't filter mid-sheet)

    wb.save(output_path)
    print(f"\n✅ Excel saved to: {output_path}")
    print(f"   {len(sorted_tcodes)} unique transaction code(s) found.")
    print(f"   {len(others)} document(s) with no codes → 'Others'.")


def main():
    # input/ and output/ folders sit next to this script — no arguments needed
    script_dir  = os.path.dirname(os.path.abspath(__file__))
    input_dir   = os.path.join(script_dir, 'input')
    output_dir  = os.path.join(script_dir, 'output')
    output_file = os.path.join(output_dir, 'SAP_Tcodes_Extract.xlsx')

    os.makedirs(input_dir,  exist_ok=True)
    os.makedirs(output_dir, exist_ok=True)

    print(f"Input  folder : {input_dir}")
    print(f"Output folder : {output_dir}\n")

    tcode_data, others = scan_folder(input_dir)
    build_excel(tcode_data, others, output_file)


if __name__ == '__main__':
    main()
