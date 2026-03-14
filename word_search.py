"""
Word Search in Word Documents
==============================
Scans all .docx files in the input/ folder (and subfolders) for one or more
search words, and produces an Excel report.

  Column A: Document Name  (relative path from input folder)
  Column B: Word Found     (the matched word from your search list)

Usage:
  python word_search.py "apple"
  python word_search.py "apple, banana, cherry"
  python word_search.py "CJ20N, FAGLB03, FB01"

Output:
  output/Word_Search_Report.xlsx

Notes:
  - Search is case-insensitive
  - Full-word match only (e.g. "park" will NOT match "parking")
  - Each unique (document, word) pair appears once in the report
  - Documents with no matches are listed under a "Not Found" section
"""

import os
import re
import sys
from collections import defaultdict

try:
    from docx import Document
except ImportError:
    raise SystemExit("python-docx not installed. Run: pip install python-docx")

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
except ImportError:
    raise SystemExit("openpyxl not installed. Run: pip install openpyxl")


# ---------------------------------------------------------------------------
# Search
# ---------------------------------------------------------------------------

def build_pattern(words: list) -> re.Pattern:
    """
    Build a single regex that matches any of the given words as full words.
    Uses explicit lookahead/lookbehind so it won't match substrings.
    """
    sorted_words = sorted(words, key=len, reverse=True)
    alts = '|'.join(re.escape(w) for w in sorted_words)
    return re.compile(
        r'(?<![A-Za-z0-9_])(' + alts + r')(?![A-Za-z0-9_])',
        re.IGNORECASE
    )


def get_text_blocks(filepath: str) -> list:
    """Extract all unique non-empty text blocks from a .docx file."""
    doc = Document(filepath)
    seen = set()
    blocks = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t and t not in seen:
            seen.add(t)
            blocks.append(t)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in seen:
                    seen.add(t)
                    blocks.append(t)

    return blocks


def search_documents(input_dir: str, words: list) -> tuple:
    """
    Scan all .docx files for the given words.
    Returns:
      hits      : {rel_path: set of matched words}
      not_found : [rel_paths with zero matches]
      all_docs  : [all rel_paths scanned]
    """
    pattern = build_pattern(words)

    all_docx = []
    for dirpath, _, filenames in os.walk(input_dir):
        for f in filenames:
            if f.lower().endswith('.docx') and not f.startswith('~$'):
                all_docx.append(os.path.join(dirpath, f))
    all_docx.sort()

    if not all_docx:
        raise SystemExit(
            f"No .docx files found in '{input_dir}' or any subfolders.\n"
            "Place your Word documents in the input/ folder and try again."
        )

    hits = {}
    not_found = []

    print(f"Searching {len(all_docx)} document(s) for: {', '.join(words)}\n")

    for filepath in all_docx:
        rel_path = os.path.relpath(filepath, input_dir)
        try:
            blocks = get_text_blocks(filepath)
        except Exception as e:
            print(f"  [WARN] Could not read '{rel_path}': {e}")
            not_found.append(rel_path)
            continue

        matched_words = set()
        for text in blocks:
            for m in pattern.finditer(text):
                matched_words.add(m.group(1))   # preserve original casing found

        if matched_words:
            hits[rel_path] = matched_words
            print(f"  + {rel_path} -> found: {', '.join(sorted(matched_words))}")
        else:
            not_found.append(rel_path)
            print(f"  - {rel_path} -> no matches")

    return hits, not_found


# ---------------------------------------------------------------------------
# Excel report
# ---------------------------------------------------------------------------

def build_report(hits: dict, not_found: list, search_words: list, output_path: str):
    wb = Workbook()
    ws = wb.active
    ws.title = "Search Results"

    # ── Styles ────────────────────────────────────────────────────────────────
    hdr_font    = Font(name='Arial', bold=True, color='FFFFFF', size=11)
    hdr_fill    = PatternFill('solid', start_color='1F4E79')
    hdr_align   = Alignment(horizontal='center', vertical='center', wrap_text=True)

    doc_font    = Font(name='Arial', size=10)
    word_font   = Font(name='Courier New', bold=True, color='1F4E79', size=10)
    none_font   = Font(name='Arial', size=10, italic=True, color='808080')
    nf_doc_font = Font(name='Arial', size=10, color='808080')

    top_align   = Alignment(horizontal='left', vertical='top', wrap_text=True)
    mid_align   = Alignment(horizontal='center', vertical='center', wrap_text=True)

    thin_side   = Side(style='thin',   color='BDD7EE')
    med_side    = Side(style='medium', color='1F4E79')
    thin_border = Border(left=thin_side, right=thin_side,
                         top=thin_side,  bottom=thin_side)
    med_border  = Border(left=med_side,  right=med_side,
                         top=med_side,   bottom=med_side)

    alt_fill    = PatternFill('solid', start_color='EBF3FB')
    base_fill   = PatternFill('solid', start_color='FFFFFF')
    nf_fill     = PatternFill('solid', start_color='F5F5F5')

    # ── Summary banner ────────────────────────────────────────────────────────
    ws.merge_cells('A1:B1')
    summary = ws['A1']
    summary.value = f"Search words: {', '.join(search_words)}   |   " \
                    f"{len(hits)} document(s) with matches   |   " \
                    f"{len(not_found)} document(s) with no matches"
    summary.font      = Font(name='Arial', bold=True, color='1F4E79', size=10)
    summary.fill      = PatternFill('solid', start_color='DEEAF1')
    summary.alignment = Alignment(horizontal='left', vertical='center',
                                  wrap_text=True)
    summary.border    = med_border
    ws.row_dimensions[1].height = 22

    # ── Column headers ────────────────────────────────────────────────────────
    for col, h in enumerate(['Document Name', 'Word Found'], 1):
        c = ws.cell(row=2, column=col, value=h)
        c.font = hdr_font; c.fill = hdr_fill
        c.alignment = hdr_align; c.border = med_border
    ws.row_dimensions[2].height = 28

    # ── Data rows ─────────────────────────────────────────────────────────────
    row = 3
    for i, (rel_path, matched_words) in enumerate(sorted(hits.items())):
        fill = alt_fill if (i % 2 == 0) else base_fill

        # One row per matched word, doc name shown only on first row (merged)
        sorted_matches = sorted(matched_words)
        start_row = row

        for word in sorted_matches:
            c = ws.cell(row=row, column=2, value=word)
            c.font = word_font; c.fill = fill
            c.alignment = top_align; c.border = thin_border
            ws.row_dimensions[row].height = 24
            row += 1

        # Document name in col A — merged if multiple words found
        c = ws.cell(row=start_row, column=1, value=rel_path)
        c.font = doc_font; c.fill = fill
        c.alignment = mid_align if len(sorted_matches) > 1 else top_align
        c.border = med_border

        if len(sorted_matches) > 1:
            ws.merge_cells(start_row=start_row, start_column=1,
                           end_row=row - 1,    end_column=1)
            c = ws.cell(row=start_row, column=1)
            c.font = doc_font; c.fill = fill
            c.alignment = mid_align; c.border = med_border

    # ── Not Found section ─────────────────────────────────────────────────────
    if not_found:
        # Separator row
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
        sep = ws.cell(row=row, column=1, value='Documents with no matches')
        sep.font      = Font(name='Arial', bold=True, color='FFFFFF', size=10)
        sep.fill      = PatternFill('solid', start_color='808080')
        sep.alignment = Alignment(horizontal='left', vertical='center')
        sep.border    = med_border
        ws.row_dimensions[row].height = 20
        row += 1

        for rel_path in sorted(not_found):
            c = ws.cell(row=row, column=1, value=rel_path)
            c.font = nf_doc_font; c.fill = nf_fill
            c.alignment = top_align; c.border = thin_border

            c = ws.cell(row=row, column=2, value='—')
            c.font = none_font; c.fill = nf_fill
            c.alignment = top_align; c.border = thin_border

            ws.row_dimensions[row].height = 22
            row += 1

    # ── Column widths & freeze ────────────────────────────────────────────────
    ws.column_dimensions['A'].width = 50
    ws.column_dimensions['B'].width = 30
    ws.freeze_panes = 'A3'

    wb.save(output_path)
    print(f"\nReport saved to: {output_path}")
    print(f"  {len(hits)} document(s) matched.")
    print(f"  {len(not_found)} document(s) had no matches.")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 2:
        raise SystemExit(
            "Usage: python word_search.py \"word1, word2, word3\"\n"
            "Example: python word_search.py \"CJ20N, FAGLB03, FB01\""
        )

    # Parse comma-separated words from argument
    raw = sys.argv[1]
    search_words = [w.strip() for w in raw.split(',') if w.strip()]

    if not search_words:
        raise SystemExit("No valid search words provided.")

    script_dir  = os.path.dirname(os.path.abspath(__file__))
    input_dir   = os.path.join(script_dir, 'input')
    output_dir  = os.path.join(script_dir, 'output')
    output_file = os.path.join(output_dir, 'Word_Search_Report.xlsx')

    os.makedirs(input_dir,  exist_ok=True)
    os.makedirs(output_dir, exist_ok=True)

    print(f"Input  folder : {input_dir}")
    print(f"Output folder : {output_dir}\n")

    hits, not_found = search_documents(input_dir, search_words)
    build_report(hits, not_found, search_words, output_file)


if __name__ == '__main__':
    main()
