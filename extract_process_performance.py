"""
Extract "Process Performance" tables from all .docx in INPUT_FOLDER.

- File-specific columns — no overlap between documents
- Col C "Filter Category" = Header Content 1 value, repeated on every row of
  the block so filtering by category shows both sub-header and data rows.
  "Table Not Found" for files with no Process Performance table.
- Fake-merge for Folder/Filename: value on every row (AutoFilter-safe),
  text invisible on rows 2+ via font = background colour
- Files with no Process Performance table → single blank row
"""

import os
from pathlib import Path
from collections import OrderedDict

from docx import Document
from docx.oxml.ns import qn

from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ── config ─────────────────────────────────────────────────────────────────
_HERE         = os.path.dirname(os.path.abspath(__file__))
INPUT_FOLDER  = os.path.join(_HERE, "Input")
OUTPUT_FOLDER = os.path.join(_HERE, "Output")
OUTPUT_FILE   = os.path.join(OUTPUT_FOLDER, "process_performance_output.xlsx")

PERF_KEYWORDS = ("process performance", "sla", "kpi")

# ── cached style objects ───────────────────────────────────────────────────
def _fill(h):      return PatternFill("solid", fgColor=h)
def _font(b, col): return Font(name="Calibri", bold=b, color=col, size=11)
def _align(h, v):  return Alignment(horizontal=h, vertical=v, wrap_text=True)
def _side(s, c):   return Side(style=s, color=c)

_NO   = Side(style=None)
_MED  = _side("medium", "4472C4")
_THIN = _side("thin",   "BFBFBF")

def _bdr(t, b, l, r): return Border(top=t, bottom=b, left=l, right=r)

THIN_ALL = _bdr(_THIN, _THIN, _THIN, _THIN)

F_HDR, F_FN, F_SH, F_VAL, F_NA = (
    _fill("1F4E79"), _fill("F2F2F2"), _fill("2E75B6"),
    _fill("FFFFFF"), _fill("F5F5F5"),
)
FT_HDR                 = _font(True,  "FFFFFF")
FT_FN_VIS, FT_FN_HID  = _font(True,  "1F4E79"), _font(True, "F2F2F2")
FT_SH, FT_VAL, FT_NA  = _font(True,  "FFFFFF"), _font(False, "000000"), _font(False, "D0D0D0")

AL_CC, AL_LC, AL_LT = _align("center","center"), _align("left","center"), _align("left","top")

# pre-build the 4 possible fn-cell borders once
_FN_BDR = {
    (ff, fl): _bdr(_MED if ff else _NO, _MED if fl else _NO, _MED, _THIN)
    for ff in (True, False) for fl in (True, False)
}

def _apply(cell, fill, font, align, border):
    cell.fill = fill; cell.font = font
    cell.alignment = align; cell.border = border

def _fn_cell(cell, value, is_first: bool, is_last: bool):
    """Folder/Filename cell — visible font on first row, invisible on the rest."""
    _apply(cell, F_FN,
           FT_FN_VIS if is_first else FT_FN_HID,
           AL_LC, _FN_BDR[(is_first, is_last)])
    cell.value = value


# ══════════════════════════════════════════════════════════════════════════
# DOCX EXTRACTION
# ══════════════════════════════════════════════════════════════════════════

def _iter_docx(root: str):
    yield from sorted(Path(root).rglob("*.docx"))


def _cell_paras(cell) -> list:
    return [p.text.strip() for p in cell.paragraphs if p.text.strip()]


def _cell_text(cell) -> str:
    return "\n".join(_cell_paras(cell))


def _is_perf_section(doc, table) -> bool:
    """True if the nearest heading above *table* matches performance keywords."""
    body = list(doc.element.body)
    try:
        tbl_idx = body.index(table._tbl)
    except ValueError:
        return False
    for elem in reversed(body[:tbl_idx]):
        if elem.tag.split("}")[-1] != "p":
            continue
        pPr   = elem.find(qn("w:pPr"))
        style = ""
        if pPr is not None:
            ps = pPr.find(qn("w:pStyle"))
            if ps is not None:
                style = (ps.get(qn("w:val")) or "").lower().replace("-", " ")
        if any(h in style for h in ("heading", "title", "toc")):
            text = "".join(
                t.text or "" for t in elem.iter() if t.tag == qn("w:t")
            ).lower()
            return any(k in text for k in PERF_KEYWORDS)
    return False


def _extract_table(table) -> tuple:
    """(base_headers, last_col_name, logical_rows) — deduplicates repeated physical rows."""
    if not table.rows:
        return [], None, []

    rows      = table.rows
    doc_hdrs  = [_cell_text(c) or f"Col{i+1}" for i, c in enumerate(rows[0].cells)]
    base_hdrs = doc_hdrs[:-1]
    last_col  = doc_hdrs[-1]

    groups = OrderedDict()
    for row in rows[1:]:
        cells     = row.cells
        base_vals = [_cell_text(c) for c in cells[:-1]]
        paras     = _cell_paras(cells[-1])
        key       = tuple(base_vals)
        if key not in groups:
            groups[key] = {"base_vals": base_vals, "last_paras": []}
        for p in paras:
            if p not in groups[key]["last_paras"]:
                groups[key]["last_paras"].append(p)

    logical_rows = []
    for g in groups.values():
        rd = dict(zip(base_hdrs, g["base_vals"]))
        rd["_lp"] = g["last_paras"]
        logical_rows.append(rd)

    return base_hdrs, last_col, logical_rows


# ══════════════════════════════════════════════════════════════════════════
# BLOCK COLLECTION
# ══════════════════════════════════════════════════════════════════════════

def _col_list(base_hdrs: list, last_col, mlp: int) -> list:
    if not last_col:
        return base_hdrs
    suffix = [f"{last_col} {i+1}" for i in range(mlp)] if mlp > 1 else [last_col]
    return base_hdrs + suffix


def collect_blocks() -> list:
    blocks = []
    for fpath in _iter_docx(INPUT_FOLDER):
        rel    = fpath.relative_to(INPUT_FOLDER)
        folder = str(rel.parent) if str(rel.parent) != "." else "(root)"
        fname  = fpath.name
        print(f"Processing: {rel}")

        stub = dict(folder=folder, filename=fname, base_hdrs=[], last_col=None,
                    rows=[], has_table=False, mlp=0, col_list=[])
        try:
            doc = Document(str(fpath))
        except Exception as e:
            print(f"  [ERROR] {e}"); blocks.append(stub); continue

        found = False
        for ti, table in enumerate(doc.tables):
            if not _is_perf_section(doc, table):
                continue
            b_hdrs, last_col, l_rows = _extract_table(table)
            if not b_hdrs:
                continue
            found = True
            mlp  = max((len(r["_lp"]) for r in l_rows), default=0)
            cols = _col_list(b_hdrs, last_col, mlp)
            print(f"  → Table #{ti+1}: {len(l_rows)} row(s), {len(cols)} col(s)")
            blocks.append({**stub, "base_hdrs": b_hdrs, "last_col": last_col,
                           "rows": l_rows, "has_table": True, "mlp": mlp, "col_list": cols})

        if not found:
            print("  [INFO] No Process Performance table — blank row.")
            blocks.append(stub)

    return blocks


# ══════════════════════════════════════════════════════════════════════════
# EXCEL WRITER
# ══════════════════════════════════════════════════════════════════════════

def write_excel(blocks: list, output_path: str):
    if not blocks:
        print("[WARN] Nothing to write."); return

    max_cc     = max((len(b["col_list"]) for b in blocks), default=1)
    # Layout: A=Folder, B=Filename, C=Filter Category, D..=content cols
    CONTENT_START = 4   # 1-based column index where content begins (col D)
    total_cols    = CONTENT_START - 1 + max_cc   # A+B+C + content

    wb = Workbook()
    ws = wb.active
    ws.title = "Process Performance"

    # Row 1 — global header
    global_hdrs = (["Folder Name", "Word Doc Name", "Filter Category"] +
                   [f"Header Content {i+1}" for i in range(max_cc)])
    for ci, lbl in enumerate(global_hdrs, 1):
        _apply(ws.cell(row=1, column=ci, value=lbl), F_HDR, FT_HDR, AL_CC, THIN_ALL)

    cur = 2

    for blk in blocks:
        folder, fname = blk["folder"], blk["filename"]
        col_list      = blk["col_list"]
        l_rows        = blk["rows"]
        b_hdrs        = blk["base_hdrs"]
        mlp, n_own    = blk["mlp"], len(col_list)

        # Filter Category value = first label in col_list (Header Content 1),
        # or "Table Not Found" when there is no table
        filter_cat = col_list[0] if col_list else "Table Not Found"

        if blk["has_table"] and l_rows:
            n_block = 1 + len(l_rows)

            # sub-header row (blue labels)
            _fn_cell(ws.cell(row=cur, column=1), folder, True, n_block == 1)
            _fn_cell(ws.cell(row=cur, column=2), fname,  True, n_block == 1)
            _apply(ws.cell(row=cur, column=3, value=filter_cat), F_SH, FT_SH, AL_LC, THIN_ALL)
            for i, lbl in enumerate(col_list):
                _apply(ws.cell(row=cur, column=CONTENT_START+i, value=lbl),
                       F_SH, FT_SH, AL_LC, THIN_ALL)
            for ci in range(CONTENT_START+n_own, total_cols+1):   # N/A zone
                _apply(ws.cell(row=cur, column=ci), F_NA, FT_NA, AL_CC, THIN_ALL)
            cur += 1

            # data rows
            for ri, lr in enumerate(l_rows):
                is_last = ri == len(l_rows) - 1
                _fn_cell(ws.cell(row=cur, column=1), folder, False, is_last)
                _fn_cell(ws.cell(row=cur, column=2), fname,  False, is_last)
                # Filter Category repeated on every data row (same blue style)
                _apply(ws.cell(row=cur, column=3, value=filter_cat),
                       F_SH, FT_SH, AL_LC, THIN_ALL)
                for i, h in enumerate(b_hdrs):
                    _apply(ws.cell(row=cur, column=CONTENT_START+i, value=lr.get(h, "")),
                           F_VAL, FT_VAL, AL_LT, THIN_ALL)
                lp, off = lr.get("_lp", []), CONTENT_START + len(b_hdrs)
                for pi in range(mlp):
                    _apply(ws.cell(row=cur, column=off+pi,
                                   value=lp[pi] if pi < len(lp) else ""),
                           F_VAL, FT_VAL, AL_LT, THIN_ALL)
                for ci in range(CONTENT_START+n_own, total_cols+1):   # N/A zone
                    _apply(ws.cell(row=cur, column=ci), F_NA, FT_NA, AL_CC, THIN_ALL)
                cur += 1

        else:
            # no-table: single blank row, Filter Category = "Table Not Found"
            _fn_cell(ws.cell(row=cur, column=1), folder, True, True)
            _fn_cell(ws.cell(row=cur, column=2), fname,  True, True)
            _apply(ws.cell(row=cur, column=3, value="Table Not Found"),
                   F_SH, FT_SH, AL_LC, THIN_ALL)
            for ci in range(CONTENT_START, total_cols+1):
                _apply(ws.cell(row=cur, column=ci), F_VAL, FT_VAL, AL_LT, THIN_ALL)
            cur += 1

    # AutoFilter + freeze
    ws.auto_filter.ref = f"A1:{get_column_letter(total_cols)}{cur-1}"
    ws.freeze_panes    = "A2"

    # Column widths — dynamic, no hardcoding
    ws.column_dimensions["A"].width = 16
    ws.column_dimensions["B"].width = 46
    ws.column_dimensions["C"].width = 28   # Filter Category
    lbl_by_pos: dict = {}
    for blk in blocks:
        for i, lbl in enumerate(blk["col_list"]):
            lbl_by_pos.setdefault(i, set()).add(lbl)
    for i in range(max_cc):
        lbls = lbl_by_pos.get(i, set())
        w    = min(max((len(l) * 1.05 for l in lbls), default=14), 50)
        ws.column_dimensions[get_column_letter(CONTENT_START+i)].width = w

    # Row heights
    ws.row_dimensions[1].height = 28
    for r in range(2, cur):
        ws.row_dimensions[r].height = 36

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    wb.save(output_path)
    print(f"\n✅  {cur-1} rows | {total_cols} cols → {output_path}")


# ══════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    write_excel(collect_blocks(), OUTPUT_FILE)
